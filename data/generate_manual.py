"""Genera el Manual de Uso de Enola Investigadora Digital en formato DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

SCREENSHOTS_DIR = Path(__file__).parent / "screenshots_manual"
OUTPUT_PATH = Path(__file__).parent / "Manual_Enola_Investigadora_Digital.docx"

# --- Palette ---
PLUM = RGBColor(0x6B, 0x4E, 0x71)
PLUM_DEEP = RGBColor(0x5B, 0x3B, 0x5C)
ROSE = RGBColor(0xC0, 0x84, 0x97)
BRASS = RGBColor(0xBF, 0xA1, 0x81)
CHARCOAL = RGBColor(0x3A, 0x31, 0x42)
CHARCOAL_LIGHT = RGBColor(0x6B, 0x5E, 0x73)
CREAM_BG = RGBColor(0xFA, 0xF6, 0xF0)


def set_cell_shading(cell, color_hex: str) -> None:
    from docx.oxml.ns import qn

    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): color_hex,
        },
    )
    shading.append(shd)


def add_styled_paragraph(
    doc,
    text,
    font_size=11,
    bold=False,
    color=CHARCOAL,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    space_after=Pt(6),
):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def add_screenshot(doc, image_path: Path, caption: str, width_inches: float = 5.8):
    if not image_path.exists():
        add_styled_paragraph(
            doc, f"[Imagen no disponible: {image_path.name}]", font_size=10, color=CHARCOAL_LIGHT
        )
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    run_cap = cap.add_run(caption)
    run_cap.font.size = Pt(9)
    run_cap.font.italic = True
    run_cap.font.color.rgb = CHARCOAL_LIGHT
    run_cap.font.name = "Calibri"


def add_section_heading(doc, level, text):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PLUM
    return h


def add_info_box(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Light Shading"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2E3E3")
    cell.text = ""
    p_title = cell.paragraphs[0]
    run_t = p_title.add_run(title)
    run_t.font.bold = True
    run_t.font.size = Pt(10)
    run_t.font.color.rgb = PLUM
    run_t.font.name = "Calibri"
    p_body = cell.add_paragraph()
    run_b = p_body.add_run(body)
    run_b.font.size = Pt(10)
    run_b.font.color.rgb = CHARCOAL
    run_b.font.name = "Calibri"
    doc.add_paragraph()  # spacer


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.bold = True
        run_b.font.size = Pt(11)
        run_b.font.color.rgb = CHARCOAL
        run_b.font.name = "Calibri"
        run_n = p.add_run(text)
        run_n.font.size = Pt(11)
        run_n.font.color.rgb = CHARCOAL
        run_n.font.name = "Calibri"
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = CHARCOAL
        run.font.name = "Calibri"


def build_document() -> Document:
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # --- Default font ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = CHARCOAL

    # ===================================================================
    # PORTADA
    # ===================================================================
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Manual de Uso")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = PLUM
    run.font.name = "Calibri"

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle_p.add_run("Enola Investigadora Digital")
    run2.font.size = Pt(24)
    run2.font.color.rgb = ROSE
    run2.font.name = "Calibri"

    doc.add_paragraph()

    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_p.paragraph_format.space_after = Pt(8)
    run3 = desc_p.add_run(
        "Plataforma de detección de violencia de género digital\n"
        "basada en RAG (Retrieval-Augmented Generation)"
    )
    run3.font.size = Pt(13)
    run3.font.color.rgb = CHARCOAL_LIGHT
    run3.font.name = "Calibri"

    doc.add_paragraph()
    doc.add_paragraph()

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = info_p.add_run("Universidad de Granada — Máster en Cultura de Paz y DDHH\nTFM 2026")
    run4.font.size = Pt(11)
    run4.font.color.rgb = CHARCOAL_LIGHT
    run4.font.name = "Calibri"

    info_p2 = doc.add_paragraph()
    info_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = info_p2.add_run(
        "Investigadora: Kimberly Michell Luna Eraso\nTutora: María del Mar García Vita"
    )
    run5.font.size = Pt(10)
    run5.font.color.rgb = CHARCOAL_LIGHT
    run5.font.name = "Calibri"

    # Add logo
    logo_path = (
        Path(__file__).parent / "src" / "ui" / "nicegui_app" / "static" / "logo-enola-new.png"
    )
    if logo_path.exists():
        doc.add_paragraph()
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_p.add_run()
        logo_run.add_picture(str(logo_path), width=Inches(3.0))

    doc.add_page_break()

    # ===================================================================
    # TABLA DE CONTENIDOS
    # ===================================================================
    add_section_heading(doc, 1, "Tabla de Contenidos")

    toc_items = [
        ("1.", "Introducción"),
        ("2.", "Acceso a la Plataforma"),
        ("3.", "Panel de Inicio (Dashboard)"),
        ("4.", "Página de Validación Humana"),
        ("5.", "Módulo de Estadística"),
        ("6.", "Módulo de IA y Confiabilidad"),
        ("7.", "Inspector de Contenido"),
        ("8.", "Base de Conocimiento"),
        ("9.", "Chat con Enola (RAG)"),
        ("10.", "Modo Oscuro"),
        ("11.", "Glosario"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run_num = p.add_run(f"{num}  ")
        run_num.font.size = Pt(12)
        run_num.font.bold = True
        run_num.font.color.rgb = PLUM
        run_num.font.name = "Calibri"
        run_title = p.add_run(title)
        run_title.font.size = Pt(12)
        run_title.font.color.rgb = CHARCOAL
        run_title.font.name = "Calibri"

    doc.add_page_break()

    # ===================================================================
    # 1. INTRODUCCIÓN
    # ===================================================================
    add_section_heading(doc, 1, "1. Introducción")

    add_styled_paragraph(
        doc,
        "Enola Investigadora Digital es una plataforma académica diseñada para "
        "detectar contenido de violencia de género en publicaciones de Facebook. "
        "Utiliza técnicas de RAG (Retrieval-Augmented Generation) combinando "
        "un modelo de lenguaje local (Ollama) con una base de conocimiento "
        "indexada en ChromaDB.",
        font_size=11,
    )

    add_styled_paragraph(
        doc,
        "La plataforma permite a investigadores y revisores analizar, validar "
        "y generar reportes estadísticos sobre la incidencia de violencia de "
        "género digital, siguiendo una taxonomía de 6 categorías principales "
        "con 19 subdimensiones.",
        font_size=11,
    )

    add_section_heading(doc, 2, "Stack Tecnológico")
    add_bullet(doc, " Python 3.12")
    add_bullet(doc, " NiceGUI (interface web)")
    add_bullet(doc, " Ollama (modelos de lenguaje locales)")
    add_bullet(doc, " ChromaDB (base de vectores para RAG)")
    add_bullet(doc, " LangChain (orquestación LLM)")
    add_bullet(doc, " SQLite (almacenamiento de análisis y feedback)")
    add_bullet(doc, " Captura de datos de Facebook")

    add_section_heading(doc, 2, "Categorías de Violencia de Género")
    add_styled_paragraph(
        doc, "La plataforma clasifica el contenido en 6 categorías principales:", font_size=11
    )
    add_bullet(doc, " Violencia simbólica", "1. ")
    add_bullet(doc, " Cosificación / Slut-shaming", "2. ")
    add_bullet(doc, " Hostilidad / Feminicidio", "3. ")
    add_bullet(doc, " Manosfera / Antifeminismo", "4. ")
    add_bullet(doc, " Desacreditación de activistas", "5. ")
    add_bullet(doc, " Falso positivo / Salvaguarda", "6. ")

    add_styled_paragraph(
        doc,
        "Además, implementa un filtro de exclusión para detectar basura digital "
        "(CÓDIGO 99) y violencia común, que se separan del análisis principal.",
        font_size=11,
    )

    doc.add_page_break()

    # ===================================================================
    # 2. ACCESO A LA PLATAFORMA
    # ===================================================================
    add_section_heading(doc, 1, "2. Acceso a la Plataforma")

    add_styled_paragraph(
        doc,
        "Para acceder a Enola Investigadora Digital, abra su navegador web "
        "y diríjase a la dirección del servidor (por ejemplo, http://localhost:8080). "
        "La página de inicio se carga automáticamente.",
        font_size=11,
    )

    add_section_heading(doc, 2, "Inicio de Sesión")
    add_styled_paragraph(
        doc,
        "Al acceder a páginas que requieren autenticación (como Validación, "
        "Estadística o Inspector), el sistema lo redirigirá automáticamente "
        "a la página de login.",
        font_size=11,
    )

    add_bullet(doc, " Ingrese su nombre de usuario")
    add_bullet(doc, " Ingrese su contraseña")
    add_bullet(doc, " Haga clic en «Iniciar sesión»")

    add_info_box(
        doc,
        "💡 Tip",
        "Si no tiene cuenta, contacte al administrador de la plataforma. "
        "No existe registro público.",
    )

    add_screenshot(doc, SCREENSHOTS_DIR / "02_login.png", "Figura 1: Pantalla de inicio de sesión")

    add_section_heading(doc, 2, "Roles de Usuario")
    add_styled_paragraph(doc, "La plataforma maneja dos roles principales:", font_size=11)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Rol"
    hdr[1].text = "Permisos"
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    row1 = table.rows[1].cells
    row1[0].text = "admin"
    row1[
        1
    ].text = "Acceso total: crear/bloquear usuarios, cargar conocimiento, validar, ver estadísticas"

    row2 = table.rows[2].cells
    row2[0].text = "reviewer"
    row2[1].text = "Validar análisis, ver estadísticas, usar el inspector"

    for row in table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    add_styled_paragraph(doc, "", font_size=6)  # spacer

    add_section_heading(doc, 2, "Cerrar Sesión")
    add_styled_paragraph(
        doc,
        "Haga clic en «Salir» junto a su nombre de usuario en la barra "
        "superior. Al cerrar la pestaña del navegador, la sesión se "
        "cierra automáticamente.",
        font_size=11,
    )

    doc.add_page_break()

    # ===================================================================
    # 3. PANEL DE INICIO
    # ===================================================================
    add_section_heading(doc, 1, "3. Panel de Inicio (Dashboard)")

    add_styled_paragraph(
        doc,
        "El panel de inicio es la página principal de la plataforma. "
        "Muestra un resumen general del estado del análisis de contenido, "
        "incluyendo indicadores clave de rendimiento (KPIs).",
        font_size=11,
    )

    add_section_heading(doc, 2, "Elementos del Dashboard")

    add_section_heading(doc, 3, "Banner de Bienvenida")
    add_styled_paragraph(
        doc,
        "En la parte superior se muestra el banner de Enola con la "
        "ilustración de la investigadora y una breve descripción de "
        "la plataforma.",
        font_size=11,
    )

    add_section_heading(doc, 3, "Indicadores Clave (KPIs)")
    add_styled_paragraph(doc, "Los KPIs muestran información resumida del análisis:", font_size=11)
    add_bullet(doc, " Total de publicaciones analizadas")
    add_bullet(doc, " Total de comentarios analizados")
    add_bullet(doc, " Porcentaje de contenido violento detectado")
    add_bullet(doc, " Total de validaciones realizadas")

    add_section_heading(doc, 3, "Banner de Fiabilidad (Regla 1)")
    add_styled_paragraph(
        doc,
        "Se muestra un banner con el nivel de fiabilidad del análisis, "
        "indicando la proporción de basura digital detectada:",
        font_size=11,
    )
    add_bullet(doc, " Nivel OK: menos del 5% de basura digital", "🟢 ")
    add_bullet(doc, " Nivel Preventivo: entre 5% y 10%", "🟡 ")
    add_bullet(doc, " Nivel Crítico: más del 10%", "🔴 ")

    add_section_heading(doc, 3, "Chat con Enola")
    add_styled_paragraph(
        doc,
        "En la parte inferior del dashboard se encuentra el chat interactivo "
        "con Enola, que responde preguntas sobre la taxonomía y el análisis "
        "utilizando RAG.",
        font_size=11,
    )

    add_screenshot(
        doc, SCREENSHOTS_DIR / "01_inicio.png", "Figura 2: Panel de inicio con KPIs y chat"
    )

    doc.add_page_break()

    # ===================================================================
    # 4. VALIDACIÓN HUMANA
    # ===================================================================
    add_section_heading(doc, 1, "4. Página de Validación Humana")

    add_styled_paragraph(
        doc,
        "La validación humana es una de las funcionalidades centrales de "
        "la plataforma. Permite a los revisores corroborar o corregir los "
        "análisis realizados por la IA.",
        font_size=11,
    )

    add_section_heading(doc, 2, "Panel de KPIs")
    add_styled_paragraph(doc, "En la parte superior se muestra un resumen:", font_size=11)
    add_bullet(doc, " Análisis pendientes de revisión")
    add_bullet(doc, " Total de revisiones realizadas")
    add_bullet(doc, " Porcentaje de acuerdo con la IA")
    add_bullet(doc, " Correcciones realizadas")
    add_bullet(doc, " Correcciones indexadas en ChromaDB")

    add_section_heading(doc, 2, "Filtros")
    add_styled_paragraph(doc, "Puede filtrar los análisis por:", font_size=11)
    add_bullet(doc, " Tipo: posts, comentarios, o todos")
    add_bullet(doc, " Estado: todos, pendientes, acuerdo, corregidos")
    add_bullet(doc, " Solo violentos: mostrar únicamente contenido clasificado como violento")

    add_section_heading(doc, 2, "Revisión de Análisis")
    add_styled_paragraph(doc, "Cada análisis muestra:", font_size=11)
    add_bullet(doc, " El texto original del post o comentario")
    add_bullet(doc, " La clasificación de la IA (categoría, dimensión, severidad)")
    add_bullet(doc, " La justificación y evidencia de la clasificación")
    add_bullet(doc, " Los marcadores lingüísticos detectados")

    add_section_heading(doc, 3, "Formulario Multi-etiqueta")
    add_styled_paragraph(
        doc,
        "Un mismo contenido puede carry hasta 5 etiquetas de clasificación. El formulario permite:",
        font_size=11,
    )
    add_bullet(doc, " Agree (estoy de acuerdo con la IA)")
    add_bullet(doc, " Disagree (no estoy de acuerdo)")
    add_bullet(doc, " Agregar/quitar etiquetas")
    add_bullet(doc, " Modificar categoría, dimensión, severidad")
    add_bullet(doc, " Agregar justificación y evidencia")

    add_section_heading(doc, 3, "Opciones de Guardado")
    add_bullet(doc, " Guardar: solo guarda en SQLite")
    add_bullet(
        doc,
        " Guardar e indexar: guarda en SQLite y envía las correcciones a ChromaDB para mejorar futuros análisis",
    )

    add_info_box(
        doc,
        "🔄 Retroalimentación",
        "Las correcciones indexadas en ChromaDB se inyectan como ejemplos "
        "few-shot en el prompt del LLM, mejorando progresivamente la "
        "calidad de las clasificaciones.",
    )

    add_screenshot(
        doc, SCREENSHOTS_DIR / "02_validacion.png", "Figura 3: Página de validación humana"
    )

    doc.add_page_break()

    # ===================================================================
    # 5. ESTADÍSTICA
    # ===================================================================
    add_section_heading(doc, 1, "5. Módulo de Estadística")

    add_styled_paragraph(
        doc,
        "El módulo de estadística genera reportes basados en las reglas "
        "metodológicas del protocolo de investigación.",
        font_size=11,
    )

    add_section_heading(doc, 2, "Regla 2: Distribución de Frecuencias")
    add_styled_paragraph(doc, "Presenta una tabla con exactamente 4 columnas:", font_size=11)
    add_bullet(doc, " Categoría")
    add_bullet(doc, " Frecuencia Absoluta")
    add_bullet(doc, " Porcentaje Válido (excluye basura digital y violencia común)")
    add_bullet(doc, " Porcentaje Acumulado (llega al 100%)")

    add_section_heading(doc, 2, "Regla 3: Moda")
    add_styled_paragraph(
        doc,
        "Identifica la categoría con mayor frecuencia. Detecta "
        "distribuciones bimodales o multimodales cuando dos o más "
        "categorías comparten la frecuencia máxima.",
        font_size=11,
    )

    add_section_heading(doc, 2, "Regla 4: Análisis Bivariado")
    add_styled_paragraph(doc, "Tablas de contingencia cruzando categoría contra:", font_size=11)
    add_bullet(doc, " Subdimensión")
    add_bullet(doc, " Página de Facebook")
    add_bullet(doc, " Fecha")

    add_section_heading(doc, 2, "Gráficos")
    add_bullet(doc, " Gráfico de torta (distribución por categoría)")
    add_bullet(doc, " Gráfico de barras (ordenado de mayor a menor)")
    add_bullet(doc, " Tabla de frecuencias debajo de los gráficos")

    add_screenshot(doc, SCREENSHOTS_DIR / "03_estadistica.png", "Figura 4: Módulo de estadística")

    doc.add_page_break()

    # ===================================================================
    # 6. IA Y CONFIABILIDAD
    # ===================================================================
    add_section_heading(doc, 1, "6. Módulo de IA y Confiabilidad")

    add_styled_paragraph(
        doc,
        "Este módulo presenta las métricas de rendimiento del sistema "
        "de clasificación automática (Regla 6).",
        font_size=11,
    )

    add_section_heading(doc, 2, "Métricas Principales")
    add_bullet(doc, " Precisión: proporción de predicciones correctas sobre el total")
    add_bullet(doc, " Sensibilidad (Recall): proporción de positivos reales detectados")
    add_bullet(doc, " F1-Score: media armónica entre precisión y sensibilidad")

    add_section_heading(doc, 2, "Matriz de Confusión")
    add_styled_paragraph(doc, "Presenta los valores:", font_size=11)
    add_bullet(doc, " Verdaderos Positivos (VP): contenido violento correctamente detectado")
    add_bullet(doc, " Verdaderos Negativos (VN): contenido no violento correctamente identificado")
    add_bullet(doc, " Falsos Positivos (FP): contenido no violento marcado como violento")
    add_bullet(doc, " Falsos Negativos (FN): contenido violento no detectado")

    add_section_heading(doc, 2, "Interpretación de Métricas")
    add_styled_paragraph(
        doc,
        "Las métricas se calculan a partir del feedback de los revisores. "
        "Cada vez que un revisor valida o corrige un análisis, las métricas "
        "se actualizan automáticamente.",
        font_size=11,
    )

    add_info_box(
        doc,
        "📊 Dato importante",
        "Las métricas se calculan usando scikit-learn y se actualizan "
        "en tiempo real a medida que se reciben validaciones.",
    )

    add_screenshot(doc, SCREENSHOTS_DIR / "04_ia.png", "Figura 5: Módulo de IA y Confiabilidad")

    doc.add_page_break()

    # ===================================================================
    # 7. INSPECTOR
    # ===================================================================
    add_section_heading(doc, 1, "7. Inspector de Contenido")

    add_styled_paragraph(
        doc,
        "El inspector permite explorar y buscar en el contenido analizado. "
        "Es una herramienta para investigadores que necesitan examinar "
        "publicaciones específicas.",
        font_size=11,
    )

    add_section_heading(doc, 2, "Funcionalidades")
    add_bullet(doc, " Buscar contenido por texto")
    add_bullet(doc, " Filtrar por categoría de violencia")
    add_bullet(doc, " Filtrar por página de Facebook")
    add_bullet(doc, " Filtrar por rango de fechas")
    add_bullet(doc, " Ver el análisis completo de cada publicación")
    add_bullet(doc, " Navegar entre resultados con paginación")

    add_screenshot(doc, SCREENSHOTS_DIR / "05_inspector.png", "Figura 6: Inspector de contenido")

    doc.add_page_break()

    # ===================================================================
    # 8. BASE DE CONOCIMIENTO
    # ===================================================================
    add_section_heading(doc, 1, "8. Base de Conocimiento")

    add_styled_paragraph(
        doc,
        "La base de conocimiento es el repositorio de información que "
        "alimenta al sistema RAG. Incluye la taxonomía, glosarios y "
        "documentos de referencia.",
        font_size=11,
    )

    add_section_heading(doc, 2, "Explorar Conocimiento")
    add_styled_paragraph(
        doc,
        "Permite navegar por la estructura de carpetas del conocimiento indexado en ChromaDB:",
        font_size=11,
    )
    add_bullet(doc, " Taxonomía de categorías y subdimensiones")
    add_bullet(doc, " Glosarios de términos")
    add_bullet(doc, " Documentos de referencia metodológica")
    add_bullet(doc, " Reglas de clasificación")

    add_section_heading(doc, 2, "Cargar Documentos (Admin)")
    add_styled_paragraph(
        doc,
        "Los administradores pueden cargar nuevos documentos a ChromaDB "
        "para enriquecer la base de conocimiento. Los documentos se "
        "dividen en fragmentos (chunks) y se indexan vectorialmente.",
        font_size=11,
    )

    add_section_heading(doc, 2, "Editor de Markdown (Admin)")
    add_styled_paragraph(
        doc,
        "Permite editar directamente los archivos markdown de la "
        "taxonomía y glosarios desde la interfaz web, sin necesidad "
        "de acceder al servidor.",
        font_size=11,
    )

    add_screenshot(doc, SCREENSHOTS_DIR / "06_conocimiento.png", "Figura 7: Base de conocimiento")

    doc.add_page_break()

    # ===================================================================
    # 9. CHAT CON ENOLA
    # ===================================================================
    add_section_heading(doc, 1, "9. Chat con Enola (RAG)")

    add_styled_paragraph(
        doc,
        "El chat interactivo es una de las funcionalidades más potentes "
        "de la plataforma. Utiliza RAG para responder preguntas basándose "
        "en la base de conocimiento indexada.",
        font_size=11,
    )

    add_section_heading(doc, 2, "Cómo Usar el Chat")
    add_bullet(doc, " Escriba su pregunta en el campo de texto")
    add_bullet(doc, " Presione Enter o haga clic en el botón de enviar")
    add_bullet(doc, " Enola responderá basándose en la taxonomía y documentos indexados")
    add_bullet(doc, " Las respuestas incluyen referencias a las fuentes utilizadas")

    add_section_heading(doc, 2, "Ejemplos de Preguntas")
    add_bullet(doc, " ¿Qué es la violencia simbólica?")
    add_bullet(doc, " ¿Cuáles son las subdimensiones de la categoría 3?")
    add_bullet(doc, " ¿Cómo se clasifica el slut-shaming?")
    add_bullet(doc, " ¿Qué marcadores lingüísticos indican hostilidad?")

    add_info_box(
        doc,
        "🤖 Sobre el modelo",
        "El chat utiliza Ollama con el modelo qwen3.5:9b ejecutándose "
        "localmente. Las respuestas se generan en español argentino "
        "para mantener coherencia con el dominio de estudio.",
    )

    add_screenshot(doc, SCREENSHOTS_DIR / "07_chat.png", "Figura 8: Chat interactivo con Enola")

    doc.add_page_break()

    # ===================================================================
    # 10. MODO OSCURO
    # ===================================================================
    add_section_heading(doc, 1, "10. Modo Oscuro")

    add_styled_paragraph(
        doc,
        "La plataforma incluye un modo oscuro que se puede activar "
        "desde el interruptor en la barra superior.",
        font_size=11,
    )

    add_section_heading(doc, 2, "Activar Modo Oscuro")
    add_bullet(doc, " Localice el interruptor «Modo oscuro» en la barra superior derecha")
    add_bullet(doc, " Haga clic en el interruptor para alternar entre modo claro y oscuro")
    add_bullet(doc, " La preferencia se guarda automáticamente para su sesión")

    add_styled_paragraph(
        doc,
        "El modo oscuro adapta todos los elementos de la interfaz: "
        "fondos, textos, gráficos, tablas y el menú lateral.",
        font_size=11,
    )

    doc.add_page_break()

    # ===================================================================
    # 11. GLOSARIO
    # ===================================================================
    add_section_heading(doc, 1, "11. Glosario")

    glossary = [
        (
            "RAG",
            "Retrieval-Augmented Generation — técnica que combina recuperación "
            "de información con generación de texto por LLM.",
        ),
        ("LLM", "Large Language Model — modelo de lenguaje grande."),
        ("Ollama", "Plataforma para ejecutar modelos de lenguaje localmente."),
        ("ChromaDB", "Base de datos vectorial para almacenar embeddings de texto."),
        ("KPI", "Key Performance Indicator — indicador clave de rendimiento."),
        (
            "FPP",
            "Falso Positivo Potencial — contenido que la IA marcó como violento "
            "pero que podría no serlo.",
        ),
        (
            "CÓDIGO 99",
            "Etiqueta de exclusión para basura digital (vacíos, enlaces "
            "huérfanos, ruido tipográfico, risas, reacciones cortas).",
        ),
        (
            "Taxonomía",
            "Sistema de clasificación de 6 categorías y 19 subdimensiones "
            "para violencia de género digital.",
        ),
        ("Few-shot", "Técnica de prompt que incluye ejemplos para guiar al LLM."),
        ("Severidad", "Nivel de gravedad de la violencia detectada (baja, media, alta)."),
        ("Subdimensión", "División específica dentro de cada categoría principal."),
        (
            "Basura digital",
            "Contenido sin valor semántico: stickers, GIFs, enlaces "
            "sueltos, muletillas, monosílabos.",
        ),
        (
            "Fiabilidad",
            "Nivel de calidad del análisis basado en la proporción de basura digital detectada.",
        ),
    ]

    table = doc.add_table(rows=len(glossary) + 1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Término"
    hdr[1].text = "Definición"
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for i, (term, definition) in enumerate(glossary, 1):
        row = table.rows[i].cells
        row[0].text = term
        row[1].text = definition
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(5.0)

    # ===================================================================
    # PIE DE DOCUMENTO
    # ===================================================================
    doc.add_paragraph()
    doc.add_paragraph()

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer_p.add_run(
        "Manual de Uso — Enola Investigadora Digital\n"
        "Universidad de Granada — TFM 2026\n"
        "Generado automáticamente desde la plataforma"
    )
    run_f.font.size = Pt(9)
    run_f.font.color.rgb = CHARCOAL_LIGHT
    run_f.font.name = "Calibri"

    return doc


def main() -> None:
    doc = build_document()
    doc.save(str(OUTPUT_PATH))
    print(f"✅ Manual generado: {OUTPUT_PATH}")
    print(f"   Tamaño: {OUTPUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
