#!/usr/bin/env python3
"""Genera el informe estadístico completo en .docx con gráficos embebidos.

Lee el snapshot JSON generado por extract_stats.py y produce un documento
con las 6 reglas metodológicas, tablas formateadas y gráficos matplotlib.

Uso:
    python scripts/generate_informe_completo.py
    python scripts/generate_informe_completo.py --stats data/exports/stats_snapshot.json --out docs/informe-estadistico.docx
"""

from __future__ import annotations

import argparse
import io
import json
import textwrap
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parent.parent

# ---------------------------------------------------------------------------
# Colores del proyecto (consistente con src/ui/utils.py)
# ---------------------------------------------------------------------------
CATEGORIA_COLORS = {
    "VDG_VIOLENCIA_SIMBOLICA": "#E74C3C",
    "VDG_COSIFICACION_SLUTSHAMING": "#E67E22",
    "VDG_HOSTILIDAD_FEMINICIDIO": "#8E44AD",
    "VDG_MANOSFERA_ANTIFEMINISMO": "#2980B9",
    "VDG_DESACREDITACION_ACTIVISTAS": "#27AE60",
    "VDG_SALVAGUARDA_FALSO_POSITIVO": "#95A5A6",
}

DARK = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GRAY = RGBColor(0x7F, 0x8C, 0x8D)

# ---------------------------------------------------------------------------
# Helpers de matplotlib
# ---------------------------------------------------------------------------


def _make_pie_chart_4slices(violent: int, non_violent: int, basura: int, comun: int) -> io.BytesIO:
    """Donut chart tricotómico: violencia, no violencia, basura digital, violencia común."""
    labels = [
        "Con violencia de género",
        "Sin violencia",
        "Basura digital\n(CÓDIGO 99)",
        "Violencia común\n(sin sesgo)",
    ]
    sizes = [violent, non_violent, basura, comun]
    colors = ["#C0392B", "#2ECC71", "#F39C12", "#8E44AD"]
    explode = (0.04, 0, 0.02, 0.02)

    fig, ax = plt.subplots(figsize=(6, 4.5))
    wedges, texts, autotexts = ax.pie(
        sizes,
        labels=labels,
        autopct="%1.1f%%",
        startangle=140,
        colors=colors,
        explode=explode,
        pctdistance=0.78,
        textprops={"fontsize": 9},
    )
    for t in autotexts:
        t.set_fontsize(8)
        t.set_fontweight("bold")
    centre_circle = plt.Circle((0, 0), 0.50, fc="white")
    fig.gca().add_artist(centre_circle)
    total = sum(sizes)
    ax.text(
        0,
        0.05,
        str(total),
        ha="center",
        va="center",
        fontsize=20,
        fontweight="bold",
        color="#2C3E50",
    )
    ax.text(0, -0.12, "contenidos", ha="center", va="center", fontsize=9, color="#7F8C8D")
    ax.set_title(
        "Distribución completa del dataset (todas las categorías)",
        fontsize=10,
        fontweight="bold",
        pad=15,
    )
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def _make_bar_chart(bar_data: list[dict]) -> io.BytesIO:
    """Genera gráfico de barras horizontales por categoría."""
    cats = [r["Categoría"] for r in bar_data]
    vals = [r["Cantidad"] for r in bar_data]
    codes = [r["Código"] for r in bar_data]
    colors = [CATEGORIA_COLORS.get(c, "#7F8C8D") for c in codes]

    fig, ax = plt.subplots(figsize=(7, 4))
    y_pos = np.arange(len(cats))
    bars = ax.barh(y_pos, vals, color=colors, height=0.6, edgecolor="white", linewidth=0.5)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(cats, fontsize=9)
    ax.invert_yaxis()
    ax.set_xlabel("Cantidad de casos", fontsize=10)
    ax.set_title(
        "Distribución de categorías de violencia de género", fontsize=11, fontweight="bold", pad=12
    )

    for bar, val in zip(bars, vals):
        ax.text(
            bar.get_width() + 0.15,
            bar.get_y() + bar.get_height() / 2,
            str(val),
            va="center",
            fontsize=9,
            fontweight="bold",
        )

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.xaxis.set_major_locator(mticker.MaxNLocator(integer=True))
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def _make_confusion_heatmap(cm: dict) -> io.BytesIO:
    """Genera heatmap de la matriz de confusión."""
    matrix = np.array([[cm["VP"], cm["FP"]], [cm["FN"], cm["VN"]]])
    labels = np.array(
        [
            [f"VP\n{cm['VP']}", f"FP\n{cm['FP']}"],
            [f"FN\n{cm['FN']}", f"VN\n{cm['VN']}"],
        ]
    )

    fig, ax = plt.subplots(figsize=(5, 4))
    cmap = plt.cm.YlOrRd
    im = ax.imshow(matrix, cmap=cmap, aspect="auto")

    ax.set_xticks([0, 1])
    ax.set_xticklabels(["Pred: Violencia", "Pred: No Violencia"], fontsize=9)
    ax.set_yticks([0, 1])
    ax.set_yticklabels(["Real: Violencia", "Real: No Violencia"], fontsize=9)

    for i in range(2):
        for j in range(2):
            color = "white" if matrix[i, j] > matrix.max() / 2 else "black"
            ax.text(
                j,
                i,
                labels[i, j],
                ha="center",
                va="center",
                fontsize=12,
                fontweight="bold",
                color=color,
            )

    ax.set_title("Matriz de Confusión", fontsize=11, fontweight="bold", pad=12)
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def _make_frequency_bar(
    freq_data: list[dict], title: str = "Frecuencia por categoría"
) -> io.BytesIO:
    """Gráfico de barras verticales para distribución de frecuencias."""
    cats = [r["Categoría"] for r in freq_data if r["Frecuencia Absoluta"] > 0]
    vals = [r["Frecuencia Absoluta"] for r in freq_data if r["Frecuencia Absoluta"] > 0]
    codes = [r["Código"] for r in freq_data if r["Frecuencia Absoluta"] > 0]
    colors = [CATEGORIA_COLORS.get(c, "#7F8C8D") for c in codes]

    fig, ax = plt.subplots(figsize=(7, 4))
    x_pos = np.arange(len(cats))
    bars = ax.bar(x_pos, vals, color=colors, width=0.6, edgecolor="white", linewidth=0.5)
    ax.set_xticks(x_pos)
    ax.set_xticklabels(cats, rotation=25, ha="right", fontsize=8)
    ax.set_ylabel("Frecuencia absoluta", fontsize=10)
    ax.set_title(title, fontsize=11, fontweight="bold", pad=12)

    for bar, val in zip(bars, vals):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + 0.15,
            str(val),
            ha="center",
            fontsize=9,
            fontweight="bold",
        )

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.yaxis.set_major_locator(mticker.MaxNLocator(integer=True))
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def _make_subdimension_bar_chart(label_dist: dict) -> io.BytesIO:
    """Gráfico de barras horizontales de subdimensiones ordenado descendente."""
    rows = label_dist["rows"]
    rows_sorted = sorted(rows, key=lambda r: -r["Cantidad"])

    labels = []
    vals = []
    colors = []
    for r in rows_sorted:
        if r["Cantidad"] == 0:
            continue
        cat_label = r["Categoría"]
        dim = r["Subdimensión"]
        labels.append(f"{cat_label} · {dim}" if dim and dim != "—" else cat_label)
        vals.append(r["Cantidad"])
        colors.append(CATEGORIA_COLORS.get(r["Código"], "#7F8C8D"))

    fig, ax = plt.subplots(figsize=(7, 5))
    y_pos = np.arange(len(labels))
    bars = ax.barh(y_pos, vals, color=colors, height=0.6, edgecolor="white", linewidth=0.5)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=8)
    ax.invert_yaxis()
    ax.set_xlabel("Cantidad de casos", fontsize=10)
    ax.set_title(
        "Distribución por subdimensión (todas las categorías)",
        fontsize=10,
        fontweight="bold",
        pad=12,
    )

    for bar, val in zip(bars, vals):
        ax.text(
            bar.get_width() + 0.1,
            bar.get_y() + bar.get_height() / 2,
            str(val),
            va="center",
            fontsize=8,
            fontweight="bold",
        )

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.xaxis.set_major_locator(mticker.MaxNLocator(integer=True))
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def _make_validation_pie(val_data: dict) -> io.BytesIO:
    """Gráfico de pastel para validación humana."""
    labels = ["Acuerdo", "Desacuerdo", "Pendientes"]
    sizes = [val_data["agreed_count"], val_data["disagreed_count"], val_data["pending_count"]]
    colors = ["#2ECC71", "#E74C3C", "#BDC3C7"]
    sizes = [s for s in sizes if s > 0]
    labels = [
        l
        for l, s in zip(
            labels,
            [val_data["agreed_count"], val_data["disagreed_count"], val_data["pending_count"]],
        )
        if s > 0
    ]
    colors = [
        c
        for c, s in zip(
            colors,
            [val_data["agreed_count"], val_data["disagreed_count"], val_data["pending_count"]],
        )
        if s > 0
    ]

    if not sizes:
        return None

    fig, ax = plt.subplots(figsize=(4.5, 3.5))
    wedges, texts, autotexts = ax.pie(
        sizes,
        labels=labels,
        autopct="%1.1f%%",
        startangle=90,
        colors=colors,
        pctdistance=0.75,
        textprops={"fontsize": 9},
    )
    for t in autotexts:
        t.set_fontsize(8)
        t.set_fontweight("bold")
    ax.set_title("Resultado de validación humana", fontsize=10, fontweight="bold", pad=10)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Helpers de DOCX
# ---------------------------------------------------------------------------


def _set_cell_shading(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def _add_table(doc: Document, headers: list[str], rows: list[list], col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(hdr_cells[i], "2C3E50")
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = str(val) if val is not None else ""
            for p in cells[c_idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(8.5)
        if r_idx % 2 == 0:
            for c in cells:
                _set_cell_shading(c, "ECF0F1")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph("")
    return table


def _heading(doc: Document, level: int, text: str):
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def _bullet(doc: Document, text: str, bold_prefix: str = ""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def _add_chart(doc: Document, chart_buf: io.BytesIO, caption: str = ""):
    """Inserta un gráfico centrado con caption opcional."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(chart_buf, width=Inches(5.5))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.size = Pt(8)
        r.italic = True
        r.font.color.rgb = GRAY


# ---------------------------------------------------------------------------
# Secciones del informe
# ---------------------------------------------------------------------------


def _build_portada(doc: Document, meta: dict):
    for _ in range(5):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INFORME ESTADÍSTICO")
    run.font.size = Pt(28)
    run.font.color.rgb = ACCENT
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Sistema de Detección de Violencia de Género Digital en Facebook")
    run.font.size = Pt(14)
    run.font.color.rgb = DARK

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Clasificación RAG con Ollama + ChromaDB + LangChain")
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    run.italic = True

    doc.add_paragraph("")
    doc.add_paragraph("")

    info_lines = [
        f"Fecha de generación: {meta['generated_at'][:10]}",
        f"Base de datos: {Path(meta['db_path']).name}",
        f"Registros totales: {meta['total_analysis']} ({meta['total_posts']} posts, {meta['total_comments']} comentarios)",
        f"Páginas de Facebook muestreadas: {meta['total_pages']}",
        f"Revisiones humanas registradas: {meta['total_feedback']}",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(10)

    doc.add_page_break()


def _build_resumen_ejecutivo(doc: Document, data: dict):
    meta = data["meta"]
    kpis = data["kpis"]
    rel = data["reliability"]
    val = data["validation_breakdown"]
    rm = data["reliability_metrics"]

    _heading(doc, 1, "Resumen Ejecutivo")

    _para(
        doc,
        (
            "El presente informe consolida los resultados estadísticos del sistema automatizado "
            "de clasificación de contenido violento por razones de género en plataformas de Facebook. "
            "La arquitectura del sistema se basa en Retrieval-Augmented Generation (RAG), utilizando "
            "Ollama como modelo de lenguaje local, ChromaDB como almacén de vectores y LangChain "
            "como orquestador del pipeline de clasificación. El conjunto de datos está compuesto por "
            f"{meta['total_analysis']} registros extraídos de {meta['total_pages']} páginas de Facebook, "
            f"de los cuales {meta['total_posts']} son publicaciones y {meta['total_comments']} son "
            "comentarios. Cada contenido puede recibir hasta 5 etiquetas simultáneas, lo que refleja "
            "la naturaleza multinivel de la violencia de género en entornos digitales."
        ),
    )

    _heading(doc, 2, "Indicadores clave del dataset")

    kpi_rows = [
        ["Total de registros analizados", str(kpis["total"])],
        ["Contenidos clasificados como violentos", f"{kpis['violent']} ({kpis['violent_pct']}%)"],
        ["Categorías taxonómicas canónicas", str(kpis["categories"])],
        ["Páginas de Facebook muestreadas", str(kpis["pages"])],
        ["Categoría con mayor frecuencia absoluta", kpis["top_category"]],
        ["Archivos del glosario de conocimiento", str(kpis["knowledge_files"])],
        [
            "Total de votos de clasificación (multi-etiqueta)",
            str(data["label_distribution"]["total_label_votes"]),
        ],
    ]
    _add_table(doc, ["Indicador", "Valor"], kpi_rows, col_widths=[3.5, 2])

    _heading(doc, 2, "Resumen de validación humana")

    val_rows = [
        [
            "Registros con validación cruzada",
            f"{val['validated_count']} de {val['total']} ({val['validated_pct']}%)",
        ],
        [
            "Casos donde IA y humano coinciden",
            f"{val['agreed_count']} ({round(val['agreed_count'] / val['validated_count'] * 100, 1) if val['validated_count'] else 0}%)",
        ],
        [
            "Casos donde el humano corrige a la IA",
            f"{val['disagreed_count']} ({round(val['disagreed_count'] / val['validated_count'] * 100, 1) if val['validated_count'] else 0}%)",
        ],
        [
            "Contenidos que requirieron ajuste manual",
            f"{data['adjustment_breakdown']['adjusted_count']} ({data['adjustment_breakdown']['adjusted_pct']}%)",
        ],
    ]
    _add_table(doc, ["Métrica", "Valor"], val_rows, col_widths=[3.5, 2])

    # Validation chart
    val_chart = _make_validation_pie(val)
    if val_chart:
        _add_chart(doc, val_chart, "Figura 1 — Distribución de resultados de la validación humana")

    # Brief mention of Regla 6 performance
    _para(
        doc,
        (
            f"En términos de rendimiento del clasificador automático, la precisión alcanzada fue del "
            f"{rm['Precisión'] * 100:.1f}% y la sensibilidad (recall) del {rm['Sensibilidad (Recall)'] * 100:.1f}%, "
            f"lo que arroja un F1-Score de {rm['F1-Score'] * 100:.1f}%. Estos valores se analizan en "
            f"profundidad en la Regla 6 de este informe."
        ),
    )

    doc.add_page_break()


def _build_regla1(doc: Document, data: dict):
    rel = data["reliability"]

    _heading(doc, 1, "Regla 1 — Fiabilidad de datos (valores perdidos)")

    _para(
        doc,
        (
            "La fiabilidad de datos constituye el primer filtro del análisis estadístico. Su propósito "
            "es evaluar la calidad de la información capturada por el scraper antes de que ingrese al "
            "pipeline de clasificación. No se trata de eliminar registros, sino de identificar y marcar "
            "aquellos que no aportan información analizable: contenidos vacíos (imágenes sin texto, "
            "stickers, GIFs), enlaces huérfanos, ruido tipográfico puro, o expresiones de risa sin "
            'carga semántica ("jajaja", "jeje", etc.). Estos registros se etiquetan con el sentinela '
            "CODIGO_99 y se excluyen de los denominadores de las Reglas 2, 3 y 4, pero permanecen en "
            "la base de datos para efectos de transparencia y auditoría."
        ),
    )

    _para(
        doc,
        (
            "El sistema también detecta violencia común sin sesgo de género, es decir, agresiones que "
            "no están dirigidas a víctimas por su condición de mujer. Estos registros se marcan con "
            "el sentinela VIOLENCIA_COMUN y se excluyen del mismo modo. La distinción es relevante: "
            "el objetivo de este estudio es mapear la violencia de género específica, no la violencia "
            "general."
        ),
    )

    _heading(doc, 2, "Tabla resumen de exclusiones")

    excl_rows = [
        ["Total de registros en la base", str(rel["total"])],
        ["Basura digital (CODIGO_99)", f"{rel['n_basura_digital']} ({rel['pct_basura']}%)"],
        [
            "Violencia común sin sesgo de género",
            f"{rel['n_violencia_comun']} ({rel['pct_violencia_comun']}%)",
        ],
        [
            "Total excluido",
            f"{rel['n_basura_digital'] + rel['n_violencia_comun']} ({round(rel['pct_basura'] + rel['pct_violencia_comun'], 2)}%)",
        ],
        [
            "Registros válidos para análisis",
            str(rel["total"] - rel["n_basura_digital"] - rel["n_violencia_comun"]),
        ],
    ]
    _add_table(doc, ["Métrica", "Valor"], excl_rows, col_widths=[3.5, 2])

    _heading(doc, 2, "Desglose por condición de exclusión")

    cond_rows = []
    for code, count in rel["detalle_basura_codigos"].items():
        pct = round(count / rel["total"] * 100, 2)
        cond_rows.append([code, str(count), f"{pct}%"])
    _add_table(
        doc,
        ["Código de condición", "Cantidad", "% del total"],
        cond_rows,
        col_widths=[2.5, 1.5, 1.5],
    )

    _heading(doc, 2, "Diagnóstico del sistema")
    _para(doc, rel["mensaje"], italic=True)

    _heading(doc, 2, "Observaciones y hallazgos")

    if rel["nivel"] == "critica":
        _bullet(
            doc,
            (
                f"Con un {rel['pct_basura']}% de basura digital, el dataset supera ampliamente el umbral "
                f"crítico del 10%. Esto implica que más de una de cada seis extracciones del scraper no "
                f"contiene información analizable. La causa inmediata es la captura de contenido multimedia "
                f"(imágenes, stickers, GIFs) cuyo texto no puede ser extraído. La consecuencia directa es "
                f"que la muestra efectiva de análisis se reduce significativamente respecto al volumen total "
                f"descargado."
            ),
            "Nivel de alerta CRÍTICO: ",
        )
    elif rel["nivel"] == "preventiva":
        _bullet(
            doc,
            (
                f"El {rel['pct_basura']}% de basura digital se encuentra en zona de precaución (5-10%). "
                f"No es alarmante, pero conviene monitorear la evolución para evitar que se degrade."
            ),
            "Nivel preventivo: ",
        )
    else:
        _bullet(
            doc,
            (
                f"El {rel['pct_basura']}% de basura digital es bajo (<5%), lo que indica una extracción "
                f"limpia y un scraper que funciona de manera adecuada."
            ),
            "Nivel aceptable: ",
        )

    all_empty = all("VACIO" in k or "NA" in k or "NAN" in k for k in rel["detalle_basura_codigos"])
    if all_empty:
        _bullet(
            doc,
            (
                "La totalidad de los registros marcados como basura digital corresponden a la condición "
                "COND_1_VACIO, lo que indica que el problema no es ruido tipográfico ni enlaces rotos, "
                "sino contenido puramente visual (fotos, stickers, GIFs) que el scraper descarga pero "
                "que no tiene texto asociado. Esto sugiere que una mejora del scraper debería incluir "
                "un filtro previo que descarte payloads vacíos antes de persistirlos en la base de datos."
            ),
            "Patrón detectado: ",
        )

    if rel["n_violencia_comun"] > 0:
        _bullet(
            doc,
            (
                f"Se detectaron {rel['n_violencia_comun']} registros ({rel['pct_violencia_comun']}%) "
                f"de violencia que no cumple con los criterios de violencia de género. Estos son "
                f"agresiones genéricas sin marcadores de género específicos. Aunque su proporción es "
                f"baja, su exclusión es metodológicamente necesaria para no contaminar los resultados."
            ),
            "Violencia común: ",
        )

    doc.add_page_break()


def _build_regla2(doc: Document, data: dict):
    freq_cat = data["frequency_categoria"]
    freq_sub = data["frequency_subdimension"]

    _heading(doc, 1, "Regla 2 — Distribución de frecuencias")

    _para(
        doc,
        (
            "La distribución de frecuencias responde a una pregunta fundamental: ¿qué tipos de violencia "
            "de género aparecen con mayor frecuencia en el corpus analizado? Para responderlo, el sistema "
            "cuenta cuántas veces aparece cada categoría taxonómica y cada subdimensión, calculando tanto "
            "la frecuencia absoluta como el porcentaje válido sobre el total de registros excluyendo "
            "basura digital y violencia común. La tabla se ordena de mayor a menor frecuencia, y se "
            "incluye un porcentaje acumulado que permite identificar rápidamente hasta qué punto las "
            "categorías principales concentran la violencia detectada."
        ),
    )

    _para(
        doc,
        (
            "Un aspecto relevante es que el sistema es multi-etiqueta: cada contenido puede recibir hasta "
            "5 etiquetas simultáneas. Esto significa que un solo post puede contribuir más de un voto a "
            "la distribución. Esta característica refleja la realidad de la violencia de género digital, "
            "donde una misma publicación puede combinar insultos (Manosfera), cosificación (Mercantilización "
            "Corporal) y minimización (Salvaguarda) en un solo mensaje."
        ),
    )

    _heading(doc, 2, "Frecuencia por categoría")

    total_validos = sum(r["Frecuencia Absoluta"] for r in freq_cat)
    cat_rows = []
    for r in freq_cat:
        cat_rows.append(
            [
                r["Categoría"],
                str(r["Frecuencia Absoluta"]),
                f"{r['Porcentaje Válido']}%",
                f"{r['Porcentaje Acumulado']}%",
            ]
        )
    cat_rows.append(["TOTAL", str(total_validos), "100%", "—"])
    _add_table(
        doc,
        ["Categoría", "Frec. absoluta", "% válido", "% acumulado"],
        cat_rows,
        col_widths=[2.5, 1.2, 1.2, 1.2],
    )

    # Bar chart
    freq_chart = _make_frequency_bar(freq_cat, "Distribución de frecuencias por categoría")
    _add_chart(
        doc,
        freq_chart,
        "Figura 2 — Distribución de frecuencias por categoría (solo registros válidos)",
    )

    _heading(doc, 2, "Frecuencia por subdimensión (todas con casos > 0)")

    sub_with_freq = [r for r in freq_sub if r["Frecuencia Absoluta"] > 0]
    sub_rows = []
    for r in sub_with_freq:
        sub_rows.append(
            [
                r["Código"],
                str(r["Frecuencia Absoluta"]),
                f"{r['Porcentaje Válido']}%",
                f"{r['Porcentaje Acumulado']}%",
            ]
        )
    _add_table(
        doc,
        ["Subdimensión", "Frec. absoluta", "% válido", "% acumulado"],
        sub_rows,
        col_widths=[1.5, 1.5, 1.5, 1.5],
    )

    _heading(doc, 2, "Observaciones")

    top_cat = freq_cat[0] if freq_cat else None
    if top_cat:
        _bullet(
            doc,
            (
                f"La categoría «{top_cat['Categoría']}» concentra el {top_cat['Porcentaje Válido']}% "
                f"de los casos válidos con {top_cat['Frecuencia Absoluta']} apariciones. Esto indica que "
                f"este tipo de violencia simbólica es el patrón predominante en las páginas muestreadas."
            ),
            "Categoría dominante: ",
        )

    top2_accum = freq_cat[1]["Porcentaje Acumulado"] if len(freq_cat) > 1 else 0
    _bullet(
        doc,
        (
            f"Las dos categorías más frecuentes acumulan el {top2_accum}% del total válido, lo cual "
            f"sugiere una distribución relativamente concentrada en un reducido número de tipologías. "
            f"Las cuatro categorías restantes reparten el {100 - top2_accum}% restante."
        ),
        "Concentración: ",
    )

    top_sub = sub_with_freq[0] if sub_with_freq else None
    if top_sub:
        _bullet(
            doc,
            (
                f"La subdimensión «{top_sub['Código']}» lidera con {top_sub['Frecuencia Absoluta']} "
                f"casos ({top_sub['Porcentaje Válido']}%). Esto permite afinar el análisis más allá "
                f"de la categoría padre y entender qué variantes específicas de violencia son más frecuentes."
            ),
            "Subdimensión líder: ",
        )

    n_zeros = len([r for r in freq_sub if r["Frecuencia Absoluta"] == 0])
    if n_zeros > 0:
        _bullet(
            doc,
            (
                f"De las 19 subdimensiones definidas en la taxonomía, {n_zeros} no presentaron ningún "
                f"caso en la muestra actual. Esto puede deberse a un tamaño de muestra reducido o a "
                f"que esas variantes son genuinamente menos frecuentes en las páginas analizadas."
            ),
            "Subdimensiones vacías: ",
        )

    doc.add_page_break()


def _build_regla3(doc: Document, data: dict):
    mode_cat = data["mode_categoria"]
    mode_sub = data["mode_subdimension"]

    _heading(doc, 1, "Regla 3 — Moda (medida de tendencia central)")

    _para(
        doc,
        (
            "La moda identifica el valor más frecuente dentro de una distribución. A diferencia de la "
            "media, que se calcula con valores numéricos, la moda funciona con datos categóricos como "
            "las tipologías de violencia de género. El sistema detecta automáticamente si la distribución "
            "es unimodal (una sola moda), bimodal (dos categorías empatadas) o multimodal (tres o más "
            "empatadas). Esta distinción es importante porque una distribución multimodal puede indicar "
            "que no hay un patrón dominante claro, sino múltiples formas de violencia que coexisten con "
            "similar frecuencia."
        ),
    )

    _heading(doc, 2, "Moda por categoría")

    modas_cat = mode_cat["modas"]
    es_multi_cat = mode_cat["es_multimodal"]
    frecuencias_cat = mode_cat["frecuencias"]

    moda_rows = []
    for code, freq in sorted(frecuencias_cat.items(), key=lambda x: -x[1]):
        es_moda = code in modas_cat
        label = f"{'★ ' if es_moda else ''}{code}"
        moda_rows.append([label, str(freq)])
    _add_table(doc, ["Categoría", "Frecuencia"], moda_rows, col_widths=[3, 1.5])

    _heading(doc, 2, "Moda por subdimensión")

    modas_sub = mode_sub["modas"]
    frecuencias_sub = mode_sub["frecuencias"]

    moda_sub_rows = []
    for code, freq in sorted(frecuencias_sub.items(), key=lambda x: -x[1]):
        es_moda = code in modas_sub
        label = f"{'★ ' if es_moda else ''}{code}"
        moda_sub_rows.append([label, str(freq)])
    _add_table(doc, ["Subdimensión", "Frecuencia"], moda_sub_rows, col_widths=[3, 1.5])

    _heading(doc, 2, "Interpretación del sistema")
    _para(doc, mode_cat["texto_descriptivo"], italic=True)

    _heading(doc, 2, "Observaciones")

    if es_multi_cat:
        _bullet(
            doc,
            (
                f"La distribución es multimodal con {len(modas_cat)} categorías empatadas en la posición "
                f"de moda. Esto sugiere que la violencia de género en las páginas analizada no se concentra "
                f"en un único tipo, sino que se manifiesta de forma diversa. Sin embargo, para sacar "
                f"conclusiones más robustas se necesitaría un tamaño de muestra mayor."
            ),
            "Distribución multimodal: ",
        )
    else:
        _bullet(
            doc,
            (
                f"La distribución es unimodal con moda en «{modas_cat[0]}». Este resultado indica que "
                f"existe un patrón dominante claro de ciberviolencia de género en la muestra. La categoría "
                f"«{modas_cat[0]}» con {frecuencias_cat[modas_cat[0]]} casos se impone sobre las demás "
                f"con una frecuencia notable."
            ),
            "Distribución unimodal: ",
        )

    if modas_sub:
        _bullet(
            doc,
            (
                f"A nivel de subdimensión, la moda es «{modas_sub[0]}» con {frecuencias_sub[modas_sub[0]]} "
                f"casos. Esto permite identificar la variante específica más prevalente, lo cual es útil "
                f"para focalizar estrategias de intervención o monitoreo."
            ),
            "Subdimensión modal: ",
        )

    doc.add_page_break()


def _build_regla4(doc: Document, data: dict):
    ct = data["crosstab_subdim"]

    _heading(doc, 1, "Regla 4 — Análisis bivariado (tabulaciones cruzadas)")

    _para(
        doc,
        (
            "El análisis bivariado cruza la categoría de violencia con dimensiones independientes para "
            "detectar patrones de asociación. Si una subdimensión aparece exclusivamente dentro de una "
            "categoría, hay una asociación clara; si se distribuye equitativamente entre varias, la "
            "relación es más difusa. El sistema presenta tres cruces: categoría por subdimensión, "
            "categoría por página de Facebook y categoría por mes de publicación."
        ),
    )

    _para(
        doc,
        (
            "Los porcentajes marginales de columna son la métrica clave de esta regla. Responden a la "
            "pregunta: «de todo el contenido que recibió la dimensión X, ¿qué porcentaje corresponde a "
            "cada categoría de violencia?». Un porcentaje del 100% en una celda indica que esa subdimensión "
            "es exclusiva de esa categoría en la muestra actual."
        ),
    )

    _heading(doc, 2, "Tabla cruzada: categoría × subdimensión")

    filas = ct["filas"]
    columnas = ct["columnas"]
    freqs = ct["frecuencias"]
    pcts = ct["porcentajes_marginales"]

    headers = ["Categoría"] + columnas
    rows_data = []
    for i, fila in enumerate(filas):
        row = [fila] + [str(freqs[i][j]) for j in range(len(columnas))]
        rows_data.append(row)
    _add_table(doc, headers, rows_data)

    _heading(doc, 2, "Porcentajes marginales de columna (%)")

    pct_headers = ["Categoría"] + columnas
    pct_rows_data = []
    for i, fila in enumerate(filas):
        row = [fila] + [f"{pcts[i][j]:.1f}%" for j in range(len(columnas))]
        pct_rows_data.append(row)
    _add_table(doc, pct_headers, pct_rows_data)

    _heading(doc, 2, "Alerta de patrón detectado")
    if ct["alerta"]:
        _para(doc, ct["alerta"], italic=True)

    # Page crosstab
    ct_pag = data["crosstab_pagina"]
    if ct_pag["columnas"]:
        _heading(doc, 2, "Tabla cruzada: categoría × página de Facebook")
        pag_headers = ["Categoría"] + ct_pag["columnas"]
        pag_rows = []
        for i, fila in enumerate(ct_pag["filas"]):
            row = [fila] + [
                str(ct_pag["frecuencias"][i][j]) for j in range(len(ct_pag["columnas"]))
            ]
            pag_rows.append(row)
        _add_table(doc, pag_headers, pag_rows)

        if ct_pag["alerta"]:
            _para(doc, ct_pag["alerta"], italic=True)

    # Date crosstab
    ct_fecha = data["crosstab_fecha"]
    if ct_fecha["columnas"] and ct_fecha["columnas"] != ["Sin fecha"]:
        _heading(doc, 2, "Tabla cruzada: categoría × mes de publicación")
        fecha_headers = ["Categoría"] + ct_fecha["columnas"]
        fecha_rows = []
        for i, fila in enumerate(ct_fecha["filas"]):
            row = [fila] + [
                str(ct_fecha["frecuencias"][i][j]) for j in range(len(ct_fecha["columnas"]))
            ]
            fecha_rows.append(row)
        _add_table(doc, fecha_headers, fecha_rows)

        if ct_fecha["alerta"]:
            _para(doc, ct_fecha["alerta"], italic=True)

    _heading(doc, 2, "Observaciones")

    if ct["alerta"]:
        _bullet(doc, ct["alerta"], "Patrón detectado: ")

    if ct_pag["alerta"]:
        _bullet(doc, ct_pag["alerta"], "Distribución por página: ")

    if ct_fecha["alerta"]:
        _bullet(doc, ct_fecha["alerta"], "Patrón temporal: ")

    _bullet(
        doc,
        (
            "Las tablas cruzadas permiten ver la estructura interna de la violencia: no basta con saber "
            "que «Violencia Simbólica» es la categoría más frecuente, sino que es relevante entender "
            "qué subdimensiones la componen y cómo se distribuye entre las diferentes páginas."
        ),
        "Utilidad del cruce: ",
    )

    doc.add_page_break()


def _build_regla5(doc: Document, data: dict):
    rel = data["reliability"]
    bar = data["bar_data"]
    label_dist = data["label_distribution"]
    kpis = data["kpis"]

    violent = kpis["violent"]
    non_violent = rel["total"] - rel["n_basura_digital"] - rel["n_violencia_comun"] - violent
    basura = rel["n_basura_digital"]
    comun = rel["n_violencia_comun"]
    total_all = rel["total"]

    _heading(doc, 1, "Regla 5 — Dashboard (visualización)")

    _para(
        doc,
        (
            "La quinta regla presenta los indicadores clave de forma visual. El dashboard incluye "
            "tres componentes: un gráfico de pastel que desglosa la composición completa del dataset "
            "(violencia de género, contenido no violento, basura digital y violencia común), un gráfico "
            "de barras con la distribución por categoría de violencia, y un gráfico de barras de "
            "subdimensiones que permite identificar qué variantes específicas son más frecuentes. "
            "Los tres gráficos utilizan la paleta de colores definida en la taxonomía para mantener "
            "coherencia visual con el resto del sistema."
        ),
    )

    # --- PIE CHART: 4 tajadas (todos los 69 registros) ---
    _heading(doc, 2, "Composición completa del dataset")

    pie_chart = _make_pie_chart_4slices(violent, non_violent, basura, comun)
    _add_chart(
        doc,
        pie_chart,
        (
            "Figura 3 — Distribución del dataset completo: violencia de género, "
            "no violencia, basura digital (CÓDIGO 99) y violencia común"
        ),
    )

    pct_v = round(violent / total_all * 100, 1) if total_all else 0
    pct_nv = round(non_violent / total_all * 100, 1) if total_all else 0
    pct_b = round(basura / total_all * 100, 1) if total_all else 0
    pct_c = round(comun / total_all * 100, 1) if total_all else 0

    pie_rows = [
        ["Con violencia de género", str(violent), f"{pct_v}%"],
        ["Sin violencia", str(non_violent), f"{pct_nv}%"],
        ["Basura digital (CÓDIGO 99)", str(basura), f"{pct_b}%"],
        ["Violencia común (sin sesgo)", str(comun), f"{pct_c}%"],
        ["TOTAL", str(total_all), "100%"],
    ]
    _add_table(doc, ["Categoría", "Cantidad", "Porcentaje"], pie_rows, col_widths=[2.5, 1.2, 1.2])

    # --- BARRAS CATEGORÍAS ---
    _heading(doc, 2, "Distribución por categoría de violencia de género")

    bar_chart = _make_bar_chart(bar)
    _add_chart(
        doc,
        bar_chart,
        (
            "Figura 4 — Distribución de categorías de violencia de género "
            "(solo contenido violento válido)"
        ),
    )

    bar_rows = []
    for r in bar:
        bar_rows.append([r["Categoría"], str(r["Cantidad"]), f"{r['Porcentaje']}%"])
    _add_table(doc, ["Categoría", "Cantidad", "Porcentaje"], bar_rows, col_widths=[2.5, 1.2, 1.2])

    # --- BARRAS SUBDIMENSIONES ---
    _heading(doc, 2, "Distribución por subdimensión")

    sub_chart = _make_subdimension_bar_chart(label_dist)
    _add_chart(
        doc,
        sub_chart,
        (
            "Figura 5 — Distribución de subdimensiones de violencia de género "
            "(todas las categorías, orden descendente)"
        ),
    )

    sub_rows = []
    for r in label_dist["rows"]:
        if r["Cantidad"] > 0:
            sub_rows.append(
                [
                    r["Categoría"],
                    r["Subdimensión"] if r["Subdimensión"] and r["Subdimensión"] != "—" else "—",
                    str(r["Cantidad"]),
                    f"{r['Porcentaje']}%",
                ]
            )
    _add_table(
        doc, ["Categoría", "Subdimensión", "Cantidad", "%"], sub_rows, col_widths=[2.2, 1, 1, 1]
    )

    # --- OBSERVACIONES ---
    _heading(doc, 2, "Observaciones")

    _bullet(
        doc,
        (
            f"De los {total_all} registros totales del corpus, {violent} ({pct_v}%) corresponden a "
            f"violencia de género, {non_violent} ({pct_nv}%) son contenido no violento, {basura} "
            f"({pct_b}%) son basura digital sin texto analizable, y {comun} ({pct_c}%) son agresiones "
            f"que no cumplen con los criterios de violencia de género. La proporción de contenido "
            f"excluido ({basura + comun} registros, {round(pct_b + pct_c, 1)}%) coincide con el "
            f"nivel de alerta CRÍTICO detectado en la Regla 1."
        ),
        "Composición general: ",
    )

    top_bar = bar[0] if bar else None
    if top_bar and len(bar) > 1:
        _bullet(
            doc,
            (
                f"«{top_bar['Categoría']}» concentra el {top_bar['Porcentaje']}% del contenido violento "
                f"con {top_bar['Cantidad']} casos, seguida por «{bar[1]['Categoría']}» ({bar[1]['Cantidad']} "
                f"casos, {bar[1]['Porcentaje']}%). Estas dos categorías suman "
                f"{top_bar['Porcentaje'] + bar[1]['Porcentaje']}% del total, lo que sugiere una "
                f"concentración significativa en un reducido número de tipologías de violencia."
            ),
            "Categorías principales: ",
        )

    top_subs = [r for r in label_dist["rows"] if r["Cantidad"] > 0][:3]
    if top_subs:
        dims_text = ", ".join(
            f"«{r['Subdimensión']}» ({r['Cantidad']} casos, {r['Porcentaje']}%)" for r in top_subs
        )
        _bullet(
            doc,
            (
                f"Las subdimensiones más frecuentes son {dims_text}. Esto permite ir más allá de la "
                f"categoría padre para entender qué variantes específicas de violencia son más prevalentes "
                f"en la muestra. Cada subdimensión representa una táctica o manifestación concreta de "
                f"ciberviolencia de género."
            ),
            "Subdimensiones líderes: ",
        )

    n_dims = len([r for r in label_dist["rows"] if r["Cantidad"] > 0])
    _bullet(
        doc,
        (
            f"Del total de {n_dims} subdimensiones con al menos un caso, la distribución es bastante "
            f"dispersa: las tres más frecuentes concentran el "
            f"{sum(r['Porcentaje'] for r in top_subs)}% del total de votos de clasificación, "
            f"lo que indica que la violencia de género se manifiesta en múltiples variantes "
            f"simultáneamente en las páginas analizadas."
        ),
        "Diversidad de variantes: ",
    )

    doc.add_page_break()


def _build_regla6(doc: Document, data: dict):
    cm = data["confusion_matrix"]
    rm = data["reliability_metrics"]
    val = data["validation_breakdown"]

    _heading(doc, 1, "Regla 6 — Métricas de fiabilidad y validez de la clasificación automática")

    # --- Explicación didáctica ampliada ---
    _para(
        doc,
        (
            "Esta regla evalúa qué tan confiable es el clasificador automático. Para ello, se comparan "
            "las predicciones del sistema contra el criterio de revisores humanos que examinaron cada "
            "contenido y decidieron si estaban o no de acuerdo con la clasificación. El resultado de "
            "esta comparación se organiza en una matriz de confusión, que es la herramienta fundamental "
            "para entender los aciertos y errores del modelo."
        ),
    )

    _heading(doc, 2, "¿Qué es la matriz de confusión?")

    _para(
        doc,
        (
            "Imaginemos una tabla de dos por dos. En las filas tenemos lo que el sistema predijo "
            "(«el contenido es violento» o «el contenido no es violento») y en las columnas tenemos "
            "lo que un ser humano determinó después de revisarlo. Las intersecciones producen cuatro "
            "escenarios posibles:"
        ),
    )

    _bullet(
        doc,
        (
            "El sistema dice «violentо» y el humano confirma que sí lo es. Es un acierto: el modelo "
            "detectó correctamente la violencia. Estos casos se denominan Verdaderos Positivos (VP)."
        ),
        "Verdadero Positivo (VP): ",
    )

    _bullet(
        doc,
        (
            "El sistema dice «no violento» y el humano coincide: el contenido efectivamente no es "
            "violento. Otro acierto, esta vez en la clase negativa. Se llama Verdadero Negativo (VN)."
        ),
        "Verdadero Negativo (VN): ",
    )

    _bullet(
        doc,
        (
            "El sistema dice «violentо» pero el humano revierte: el contenido no es violento. Es un "
            "error por exceso, un falso positivo (FP). El sistema «ve» violencia donde no la hay."
        ),
        "Falso Positivo (FP): ",
    )

    _bullet(
        doc,
        (
            "El sistema dice «no violento» pero el humano corrige: el contenido sí era violento. Es "
            "un error por omisión, un falso negativo (FN). El sistema dejó pasar violencia real. Este "
            "es el error más grave en un contexto de detección de contenido dañino."
        ),
        "Falso Negativo (FN): ",
    )

    _para(
        doc,
        (
            "La diferencia entre FP y FN tiene implicaciones prácticas importantes. Un falso positivo "
            "significa que un contenido fue marcado injustificadamente como violento, lo que puede "
            "generar una intervención innecesaria. Un falso negativo, en cambio, significa que contenido "
            "violento quedó sin detectar, lo cual es más grave porque implica que la víctima no recibe "
            "la protección o el monitoreo que correspondería."
        ),
    )

    _heading(doc, 2, "Matriz de confusión del sistema")

    cm_headers = ["", "Real: Violencia", "Real: No Violencia"]
    cm_rows = [
        ["Predicho: Violencia", str(cm["VP"]), str(cm["FP"])],
        ["Predicho: No Violencia", str(cm["FN"]), str(cm["VN"])],
    ]
    _add_table(doc, cm_headers, cm_rows, col_widths=[2, 1.5, 1.5])

    # Heatmap
    cm_chart = _make_confusion_heatmap(cm)
    _add_chart(doc, cm_chart, "Figura 5 — Visualización de la matriz de confusión")

    _heading(doc, 2, "¿Qué significan estas métricas?")

    _para(
        doc,
        (
            "A partir de la matriz de confusión se derivan tres indicadores de rendimiento que resumen "
            "la capacidad del clasificador en un solo número."
        ),
    )

    precision = rm["Precisión"]
    recall = rm["Sensibilidad (Recall)"]
    f1 = rm["F1-Score"]
    soporte = rm["Soporte"]

    _heading(doc, 3, "Precisión")
    _para(
        doc,
        (
            f"La precisión mide: «de todos los contenidos que el sistema marcó como violentos, ¿cuántos "
            f"realmente lo son?». Se calcula como VP / (VP + FP). En nuestro caso: {cm['VP']} / "
            f"({cm['VP']} + {cm['FP']}) = {precision:.4f}, es decir, {precision * 100:.1f}%. Esto "
            f"significa que cuando el sistema etiqueta un contenido como violento, acierta aproximadamente "
            f"7 de cada 10 veces. Es una precisión moderada: no es perfecta, pero tampoco es deficiente."
        ),
    )

    _heading(doc, 3, "Sensibilidad (Recall)")
    _para(
        doc,
        (
            f"La sensibilidad mide: «de todos los contenidos que realmente son violentos, ¿cuántos logra "
            f"detectar el sistema?». Se calcula como VP / (VP + FN). En nuestro caso: {cm['VP']} / "
            f"({cm['VP']} + {cm['FN']}) = {recall:.4f}, es decir, {recall * 100:.1f}%. Este es el "
            f"indicador más preocupante: el sistema solo detecta poco menos de la mitad de la violencia "
            f"real. Los {cm['FN']} falsos negativos representan contenido violento que pasó inadvertido "
            f"para el clasificador."
        ),
    )

    _heading(doc, 3, "F1-Score")
    _para(
        doc,
        (
            f"El F1-Score es la media armónica entre precisión y sensibilidad. Su valor oscila entre 0 "
            f"(peor caso) y 1 (perfecto). En nuestro caso: {f1:.4f}, es decir, {f1 * 100:.1f}%. Esta "
            f"cifra refleja el equilibrio entre ambos extremos: el sistema no es excesivamente permisivo "
            f"(que generaría muchos FP) ni excesivamente restrictivo (que generarían muchos FN). Sin "
            f"embargo, un F1-Score por debajo del 60% sugiere que hay un margen significativo de mejora "
            f"en la calibración del modelo."
        ),
    )

    _heading(doc, 3, "Soporte")
    _para(
        doc,
        (
            f"El soporte indica cuántos casos fueron evaluados para construir estas métricas. En este "
            f"caso: {soporte} revisiones humanas vinculadas a predicciones del sistema. Un soporte "
            f"mayor proporciona estimaciones más estables y representativas."
        ),
    )

    _heading(doc, 2, "Tabla resumen de métricas")

    perf_rows = [
        ["Verdaderos Positivos (VP)", str(cm["VP"]), "Aciertos en clase violenta"],
        ["Verdaderos Negativos (VN)", str(cm["VN"]), "Aciertos en clase no violenta"],
        ["Falsos Positivos (FP)", str(cm["FP"]), "Marcado incorrectamente como violento"],
        ["Falsos Negativos (FN)", str(cm["FN"]), "Violencia no detectada (error grave)"],
        ["Precisión", f"{precision * 100:.1f}%", f"{cm['VP']} / ({cm['VP']} + {cm['FP']})"],
        [
            "Sensibilidad (Recall)",
            f"{recall * 100:.1f}%",
            f"{cm['VP']} / ({cm['VP']} + {cm['FN']})",
        ],
        ["F1-Score", f"{f1 * 100:.1f}%", "Media armónica de precisión y sensibilidad"],
        ["Soporte total", str(soporte), "Revisiones humanas evaluadas"],
    ]
    _add_table(doc, ["Métrica", "Valor", "Interpretación"], perf_rows, col_widths=[2, 1.2, 3.2])

    _heading(doc, 2, "Observaciones y recomendaciones")

    _bullet(
        doc,
        (
            f"El sistema presenta un perfil conservador: su precisión ({precision * 100:.1f}%) es "
            f"mayor que su sensibilidad ({recall * 100:.1f}%). En la práctica, esto significa que "
            f"cuando el sistema clasifica algo como violento, suele tener razón; pero se le escapa "
            f"una proporción significativa de contenido violento real. Para un sistema de monitoreo "
            f"de violencia de género, esta asimetría es problemática porque implica subregistro."
        ),
        "Perfil del clasificador: ",
    )

    _bullet(
        doc,
        (
            f"Los {cm['FN']} falsos negativos son la prioridad de mejora. Cada uno representa "
            f"contenido violento que el sistema no logró identificar. Las causas posibles incluyen: "
            f"lenguaje implícito o codificado que el modelo no reconoce, violencia que se manifiesta "
            f"más en el contexto que en el texto, o prompts de clasificación que priorizan la precisión "
            f"sobre la exhaustividad."
        ),
        "Falsos negativos: ",
    )

    _bullet(
        doc,
        (
            f"Los {cm['FP']} falsos positivos, aunque menos graves, también merecen atención. "
            f"Pueden deberse a contenido que menciona violencia de género en contexto de denuncia "
            f"o análisis, y que el sistema confunde con contenido propiamente violento."
        ),
        "Falsos positivos: ",
    )

    _bullet(
        doc,
        (
            "Se recomienda reajustar los prompts de clasificación para aumentar la sensibilidad "
            "a costa de una ligera reducción de precisión. En detección de contenido dañino, es "
            "preferible generar algunos falsos positivos antes que dejar pasar violencia real. "
            "Además, continuar alimentando ChromaDB con las correcciones humanas (los 21 desacuerdos "
            "registrados) debería mejorar progresivamente el rendimiento del clasificador RAG."
        ),
        "Recomendación: ",
    )

    doc.add_page_break()


def _build_conclusiones(doc: Document, data: dict):
    kpis = data["kpis"]
    rel = data["reliability"]
    rm = data["reliability_metrics"]
    val = data["validation_breakdown"]
    adj = data["adjustment_breakdown"]

    _heading(doc, 1, "Conclusiones")

    _para(
        doc,
        (
            "El análisis estadístico de los 69 registros del corpus permite esbozar las siguientes "
            "conclusiones provisionales:"
        ),
    )

    _bullet(
        doc,
        (
            f"El dataset analizado contiene {data['meta']['total_analysis']} registros, de los cuales "
            f"el {kpis['violent_pct']}% fue clasificado como contenido violento por razones de género. "
            f"La categoría predominante es «{kpis['top_category']}», lo que sugiere que formas de "
            f"violencia simbólica (minimización, trivialización, estereotipos) son las más extendidas "
            f"en las páginas muestreadas."
        ),
        "Composición del corpus: ",
    )

    _bullet(
        doc,
        (
            f"La fiabilidad de datos es CRÍTICA: el {rel['pct_basura']}% del dataset corresponde a "
            f"basura digital, muy por encima del umbral del 10%. El scraper necesita revisión urgente "
            f"para evitar descargar contenido multimedia sin texto analizable."
        ),
        "Calidad de la extracción: ",
    )

    _bullet(
        doc,
        (
            f"El clasificador RAG alcanza una precisión del {rm['Precisión'] * 100:.1f}% pero una "
            f"sensibilidad del {rm['Sensibilidad (Recall)'] * 100:.1f}%, con un F1-Score de "
            f"{rm['F1-Score'] * 100:.1f}%. El sistema es más confiable cuando dice que algo es "
            f"violentо de cuando dice que no lo es. Esta asimetría debe corregirse priorizando la "
            f"detección de falsos negativos."
        ),
        "Rendimiento del clasificador: ",
    )

    agr_count = val["agreed_count"]
    val_count = val["validated_count"]
    agr_pct = round(agr_count / val_count * 100, 1) if val_count else 0
    _bullet(
        doc,
        (
            f"El {val['validated_pct']}% del dataset ha sido validado por humanos, con una tasa de "
            f"acuerdo del {agr_pct}%. El {adj['adjusted_pct']}% de contenidos requirió corrección "
            f"manual. Las correcciones con desacuerdo deben indexarse en ChromaDB como ejemplos "
            f"few-shot para cerrar la brecha entre la clasificación automática y el criterio humano."
        ),
        "Validación humana: ",
    )

    _heading(doc, 2, "Recomendaciones")

    _bullet(
        doc,
        (
            "Implementar un filtro previo en el scraper que descarte payloads vacíos antes de "
            "persistirlos, reduciendo así la basura digital y mejorando la eficiencia del pipeline."
        ),
        "Mejora del scraper: ",
    )

    _bullet(
        doc,
        (
            "Ajustar los prompts de clasificación para priorizar la sensibilidad sobre la precisión. "
            "En un dominio como la violencia de género, es más costoso dejar pasar contenido violento "
            "que marcar erróneamente contenido no violento."
        ),
        "Calibración del modelo: ",
    )

    _bullet(
        doc,
        (
            f"Indexar los {adj['adjusted_count']} desacuerdos registrados en ChromaDB para enriquecer "
            f"la base de ejemplos few-shot del RAGClassifier, lo que debería mejorar las clasificaciones "
            f"futuras de forma iterativa."
        ),
        "Retroalimentación RAG: ",
    )

    _bullet(
        doc,
        (
            "Continuar con el proceso de validación humana hasta alcanzar al menos el 80% de cobertura "
            "efectiva del dataset (actualmente el 100% está revisado pero con una muestra pequeña de "
            "69 registros). Con un corpus más amplio, la validación será aún más valiosa."
        ),
        "Cobertura de validación: ",
    )

    _bullet(
        doc,
        (
            "Dado que la muestra actual es reducida (69 registros de 2 páginas), se recomienda "
            "ampliar el muestreo a un mayor número de páginas y un período temporal más extenso "
            "para obtener conclusiones más generalizables."
        ),
        "Ampliación de la muestra: ",
    )


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def generate_docx(stats_path: str, out_path: str):
    stats = json.loads(Path(stats_path).read_text())

    doc = Document()

    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    _build_portada(doc, stats["meta"])
    _build_resumen_ejecutivo(doc, stats)

    _build_regla1(doc, stats)
    doc.add_page_break()

    _build_regla2(doc, stats)
    doc.add_page_break()

    _build_regla3(doc, stats)
    doc.add_page_break()

    _build_regla4(doc, stats)
    doc.add_page_break()

    _build_regla5(doc, stats)
    doc.add_page_break()

    _build_regla6(doc, stats)
    doc.add_page_break()

    _build_conclusiones(doc, stats)

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"✓ Informe generado: {out}")
    print(f"  Snapshot: {stats_path}")
    print(f"  Registros: {stats['meta']['total_analysis']}")
    print(f"  Feedback: {stats['meta']['total_feedback']}")


def main():
    parser = argparse.ArgumentParser(description="Generar informe estadístico completo .docx")
    parser.add_argument("--stats", default=str(ROOT / "data" / "exports" / "stats_snapshot.json"))
    parser.add_argument("--out", default=str(ROOT / "docs" / "informe-estadistico-completo.docx"))
    args = parser.parse_args()
    generate_docx(args.stats, args.out)


if __name__ == "__main__":
    main()
