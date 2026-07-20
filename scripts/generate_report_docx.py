#!/usr/bin/env python3
"""Generate the statistical report .docx from the stats snapshot.

Usage: python scripts/generate_report_docx.py [--stats data/exports/stats_snapshot.json]
                                                [--out docs/informe-estadistico-sistema.docx]
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

DARK = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GRAY = RGBColor(0x7F, 0x8C, 0x8D)
HEADER_BG = RGBColor(0x2C, 0x3E, 0x50)


def _set_cell_shading(cell, color_hex: str):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def _add_table(doc: Document, headers: list[str], rows: list[list], col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(hdr_cells[i], "2C3E50")
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = str(val) if val is not None else ""
            for p in cells[c_idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(8.5)
        if r_idx % 2 == 0:
            for c in cells:
                _set_cell_shading(c, "ECF0F1")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph("")
    return table


def _heading(doc: Document, level: int, text: str):
    """Add a heading."""
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, bold=False, italic=False, size=10.5):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def _bullet(doc: Document, text: str, bold_prefix: str = ""):
    """Add a bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------


def _build_portada(doc: Document, meta: dict):
    for _ in range(6):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INFORME ESTADÍSTICO DEL SISTEMA")
    run.font.size = Pt(26)
    run.font.color.rgb = ACCENT
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Detección de Violencia de Género Digital en Facebook")
    run.font.size = Pt(16)
    run.font.color.rgb = DARK

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Clasificación RAG con Ollama + ChromaDB + LangChain")
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    run.italic = True

    doc.add_paragraph("")
    doc.add_paragraph("")

    info_lines = [
        f"Fecha de generación: {meta['generated_at'][:10]}",
        f"Base de datos: {Path(meta['db_path']).name}",
        f"Total de análisis: {meta['total_analysis']}",
        f"Total de feedback humano: {meta['total_feedback']}",
        f"Registros: {meta['total_posts']} posts, {meta['total_comments']} comentarios",
        f"Páginas de Facebook: {meta['total_pages']}",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(10)

    doc.add_page_break()


def _build_resumen_ejecutivo(doc: Document, data: dict):
    meta = data["meta"]
    kpis = data["kpis"]
    rel = data["reliability"]

    _heading(doc, 1, "0. Resumen Ejecutivo")

    _para(
        doc,
        (
            "Este informe presenta el análisis estadístico completo del sistema de detección "
            "de violencia de género digital. El sistema utiliza un pipeline de Retrieval-Augmented "
            "Generation (RAG) que combina Ollama como modelo de lenguaje local con ChromaDB como "
            "base de vectores para clasificar contenido de Facebook en 6 categorías de violencia "
            "de género con un total de 19 subdimensiones."
        ),
    )

    _para(
        doc,
        (
            f"El dataset analizado contiene {meta['total_analysis']} registros "
            f"({meta['total_posts']} posts y {meta['total_comments']} comentarios) "
            f"provenientes de {meta['total_pages']} páginas de Facebook. "
            f"El sistema es multi-etiqueta: cada contenido puede portar hasta "
            f"5 etiquetas simultáneas, lo que genera "
            f"{data['label_distribution']['total_label_votes']} "
            f"votos de clasificación sobre {kpis['violent']} contenidos violentos."
        ),
    )

    _para(
        doc,
        (
            f"Se aplicaron {meta['total_feedback']} revisiones humanas (validación cruzada) que "
            f"resultaron en {data['validation_breakdown']['agreed_count']} acuerdos y "
            f"{data['validation_breakdown']['disagreed_count']} desacuerdos. "
            f"El nivel de fiabilidad de datos es {rel['nivel'].upper()} ({rel['pct_basura']}% "
            f"de valores perdidos, umbral crítico >10%)."
        ),
    )

    _heading(doc, 2, "Indicadores Clave (KPIs)")

    kpi_rows = [
        ["Total analizado", str(kpis["total"])],
        ["Contenidos con violencia", f"{kpis['violent']} ({kpis['violent_pct']}%)"],
        ["Categorías canónicas", str(kpis["categories"])],
        ["Páginas de Facebook", str(kpis["pages"])],
        ["Categoría más frecuente", kpis["top_category"]],
        ["Archivos de conocimiento", str(kpis["knowledge_files"])],
    ]
    _add_table(doc, ["Indicador", "Valor"], kpi_rows, col_widths=[3, 2])


def _build_regla1(doc: Document, data: dict):
    rel = data["reliability"]

    _heading(doc, 1, "1. Regla 1 — Fiabilidad de Datos (Valores Perdidos)")

    _para(
        doc,
        (
            "La primera regla de la metodología evalúa la integridad del dataset. "
            "Los registros excluidos no se eliminan de la base de datos, sino que se "
            "marcan con sentinels de exclusión: CODIGO_99 (basura digital) y "
            "VIOLENCIA_COMUN (violencia sin sesgo de género). Estos registros participan "
            "en el reporte de valores perdidos pero se excluyen de los denominadores "
            "de las Reglas 2 a 4."
        ),
    )

    _heading(doc, 2, "Resumen de Exclusiones")

    excl_rows = [
        ["Total de registros", str(rel["total"])],
        ["Basura digital (CODIGO_99)", f"{rel['n_basura_digital']} ({rel['pct_basura']}%)"],
        ["Violencia común", f"{rel['n_violencia_comun']} ({rel['pct_violencia_comun']}%)"],
        [
            "Total excluidos",
            f"{rel['n_basura_digital'] + rel['n_violencia_comun']} "
            f"({round(rel['pct_basura'] + rel['pct_violencia_comun'], 2)}%)",
        ],
        [
            "Registros válidos",
            f"{rel['total'] - rel['n_basura_digital'] - rel['n_violencia_comun']}",
        ],
        ["Nivel de alerta", rel["nivel"].upper()],
    ]
    _add_table(doc, ["Métrica", "Valor"], excl_rows, col_widths=[3, 2])

    _heading(doc, 2, "Desglose de Condiciones de Exclusión")

    cond_rows = []
    for code, count in rel["detalle_basura_codigos"].items():
        pct = round(count / rel["total"] * 100, 2)
        cond_rows.append([code, str(count), f"{pct}%"])
    _add_table(doc, ["Código", "Cantidad", "% del Total"], cond_rows, col_widths=[2.5, 1.5, 1.5])

    _heading(doc, 2, "Mensaje del Sistema")
    _para(doc, rel["mensaje"], italic=True)

    _heading(doc, 2, "Insights")

    if rel["nivel"] == "critica":
        _bullet(
            doc,
            (
                f"El {rel['pct_basura']}% de basura digital supera el umbral crítico del 10%. "
                f"Esto indica que el scraper captura una cantidad significativa de contenido "
                f"que no aporta información analizable (stickers, GIFs, respuestas vacías)."
            ),
            "Alerta crítica: ",
        )
    elif rel["nivel"] == "preventiva":
        _bullet(
            doc,
            (
                f"El {rel['pct_basura']}% de basura digital está en rango preventivo (5-10%). "
                f"Se recomienda monitorear la estabilidad del scraper."
            ),
            "Nivel preventivo: ",
        )
    else:
        _bullet(
            doc,
            (
                f"El {rel['pct_basura']}% de basura digital es aceptable (<5%). "
                f"La calidad de extracción del scraper es adecuada."
            ),
            "Nivel OK: ",
        )

    all_empty = all("VACIO" in k or "NA" in k or "NAN" in k for k in rel["detalle_basura_codigos"])
    if all_empty:
        _bullet(
            doc,
            (
                "La totalidad de la basura digital corresponde a payloads vacíos "
                "(COND_1_VACIO), lo que sugiere que el scraper descarga contenido "
                "multimedia (imágenes, stickers, GIFs) sin texto extraíble."
            ),
            "Patrón detectado: ",
        )

    if rel["n_violencia_comun"] > 0:
        _bullet(
            doc,
            (
                f"Se detectaron {rel['n_violencia_comun']} registros de violencia común "
                f"sin sesgo de género ({rel['pct_violencia_comun']}%). Estos se excluyen "
                f"por no pertenecer al dominio de estudio."
            ),
            "Violencia común: ",
        )


def _build_regla2(doc: Document, data: dict):
    freq_cat = data["frequency_categoria"]
    freq_sub = data["frequency_subdimension"]

    _heading(doc, 1, "2. Regla 2 — Distribución de Frecuencias")

    _para(
        doc,
        (
            "La segunda regla calcula la distribución de frecuencias de las categorías "
            "de violencia de género. Se presentan dos niveles de análisis: categorías "
            "(6 clases canónicas) y subdimensiones (19 variantes). Los porcentajes se "
            "calculan sobre el total de registros válidos (excluyendo CODIGO_99 y "
            "VIOLENCIA_COMUN). El sistema es multi-etiqueta: cada contenido puede "
            "contribuir más de un voto."
        ),
    )

    _heading(doc, 2, "Frecuencia por Categoría")

    total_validos = sum(r["Frecuencia Absoluta"] for r in freq_cat)
    cat_rows = []
    for r in freq_cat:
        cat_rows.append(
            [
                r["Categoría"],
                str(r["Frecuencia Absoluta"]),
                f"{r['Porcentaje Válido']}%",
                f"{r['Porcentaje Acumulado']}%",
            ]
        )
    cat_rows.append(["TOTAL", str(total_validos), "100%", "—"])

    _add_table(
        doc,
        ["Categoría", "Frec. Absoluta", "% Válido", "% Acumulado"],
        cat_rows,
        col_widths=[2.5, 1.2, 1.2, 1.2],
    )

    _heading(doc, 2, "Frecuencia por Subdimensión (Top 10)")

    sub_with_freq = [r for r in freq_sub if r["Frecuencia Absoluta"] > 0]
    sub_rows = []
    for r in sub_with_freq[:10]:
        sub_rows.append(
            [
                r["Código"],
                str(r["Frecuencia Absoluta"]),
                f"{r['Porcentaje Válido']}%",
                f"{r['Porcentaje Acumulado']}%",
            ]
        )
    _add_table(
        doc,
        ["Subdimensión", "Frec. Absoluta", "% Válido", "% Acumulado"],
        sub_rows,
        col_widths=[1.5, 1.5, 1.5, 1.5],
    )

    _heading(doc, 2, "Insights")

    top_cat = freq_cat[0] if freq_cat else None
    if top_cat:
        _bullet(
            doc,
            (
                f"La categoría más frecuente es '{top_cat['Categoría']}' con "
                f"{top_cat['Frecuencia Absoluta']} casos ({top_cat['Porcentaje Válido']}% del "
                f"total válido), acumulando el {top_cat['Porcentaje Acumulado']}% en la tabla "
                f"ordenada."
            ),
            "Categoría dominante: ",
        )

    top2_accum = freq_cat[1]["Porcentaje Acumulado"] if len(freq_cat) > 1 else 0
    _bullet(
        doc,
        (
            f"Las dos primeras categorías concentran el {top2_accum}% del total válido, "
            f"lo que indica una distribución moderadamente concentrada."
        ),
        "Concentración: ",
    )

    top_sub = sub_with_freq[0] if sub_with_freq else None
    if top_sub:
        _bullet(
            doc,
            (
                f"La subdimensión más frecuente es '{top_sub['Código']}' con "
                f"{top_sub['Frecuencia Absoluta']} casos ({top_sub['Porcentaje Válido']}%), "
                f"lo que permite identificar patrones específicos dentro de la categoría padre."
            ),
            "Subdimensión líder: ",
        )


def _build_regla3(doc: Document, data: dict):
    mode_cat = data["mode_categoria"]
    mode_sub = data["mode_subdimension"]

    _heading(doc, 1, "3. Regla 3 — Moda")

    _para(
        doc,
        (
            "La moda es el valor que aparece con mayor frecuencia en una distribución. "
            "El sistema detecta automáticamente distribuciones unimodales (una moda), "
            "bimodales (dos modas empatadas) y multimodales (tres o más modas empatadas). "
            "Se presenta tanto a nivel de categoría como de subdimensión."
        ),
    )

    _heading(doc, 2, "Moda por Categoría")

    modas_cat = mode_cat["modas"]
    es_multi_cat = mode_cat["es_multimodal"]
    frecuencias_cat = mode_cat["frecuencias"]

    moda_rows = []
    for code, freq in sorted(frecuencias_cat.items(), key=lambda x: -x[1]):
        es_moda = code in modas_cat
        label = f"{'★ ' if es_moda else ''}{code}"
        moda_rows.append([label, str(freq)])
    _add_table(doc, ["Categoría", "Frecuencia"], moda_rows, col_widths=[3, 1.5])

    _heading(doc, 2, "Moda por Subdimensión")

    modas_sub = mode_sub["modas"]
    frecuencias_sub = mode_sub["frecuencias"]

    moda_sub_rows = []
    for code, freq in sorted(frecuencias_sub.items(), key=lambda x: -x[1]):
        es_moda = code in modas_sub
        label = f"{'★ ' if es_moda else ''}{code}"
        moda_sub_rows.append([label, str(freq)])
    _add_table(doc, ["Subdimensión", "Frecuencia"], moda_sub_rows, col_widths=[3, 1.5])

    _heading(doc, 2, "Texto Descriptivo del Sistema")
    _para(doc, mode_cat["texto_descriptivo"], italic=True)

    _heading(doc, 2, "Insights")

    if es_multi_cat:
        _bullet(
            doc,
            (
                f"La distribución es multimodal con {len(modas_cat)} categorías empatadas. "
                f"Esto sugiere que no hay un patrón dominante claro y que la violencia de "
                f"género digital se manifiesta de forma diversa en la muestra."
            ),
            "Distribución multimodal: ",
        )
    else:
        _bullet(
            doc,
            (
                f"La distribución es unimodal con moda en '{modas_cat[0]}'. "
                f"Existe un patrón dominante claro de ciberviolencia de género."
            ),
            "Distribución unimodal: ",
        )

    if len(modas_sub) > 0:
        _bullet(
            doc,
            (
                f"A nivel de subdimensión, la moda es '{modas_sub[0]}' con "
                f"{frecuencias_sub[modas_sub[0]]} casos. Esto permite identificar "
                f"la variante específica más prevalente de violencia."
            ),
            "Subdimensión modal: ",
        )


def _build_regla4(doc: Document, data: dict):
    ct = data["crosstab_subdim"]

    _heading(doc, 1, "4. Regla 4 — Análisis Bivariado (Tabulación Cruzada)")

    _para(
        doc,
        (
            "La cuarta regla cruza la categoría de violencia contra dimensiones "
            "independientes para detectar patrones de asociación. Se presentan "
            "tres cruces: categoría × subdimensión, categoría × página de Facebook "
            "y categoría × mes de publicación. Los porcentajes marginales de columna "
            "responden a la pregunta: '¿Del total de contenido que recibió la dimensión X, "
            "qué porcentaje corresponde a cada categoría de violencia?'"
        ),
    )

    _heading(doc, 2, "Tabla Cruzada: Categoría × Subdimensión")

    filas = ct["filas"]
    columnas = ct["columnas"]
    freqs = ct["frecuencias"]
    pcts = ct["porcentajes_marginales"]

    # Build frequency table
    headers = ["Categoría"] + columnas
    rows_data = []
    for i, fila in enumerate(filas):
        row = [fila] + [str(freqs[i][j]) for j in range(len(columnas))]
        rows_data.append(row)
    _add_table(doc, headers, rows_data)

    _heading(doc, 2, "Porcentajes Marginales de Columna (%)")

    pct_headers = ["Categoría"] + columnas
    pct_rows_data = []
    for i, fila in enumerate(filas):
        row = [fila] + [f"{pcts[i][j]:.1f}%" for j in range(len(columnas))]
        pct_rows_data.append(row)
    _add_table(doc, pct_headers, pct_rows_data)

    _heading(doc, 2, "Alerta de Patrón Detectado")
    if ct["alerta"]:
        _para(doc, ct["alerta"], italic=True)

    # Crosstab by page
    ct_pag = data["crosstab_pagina"]
    if ct_pag["columnas"]:
        _heading(doc, 2, "Tabla Cruzada: Categoría × Página")

        pag_headers = ["Categoría"] + ct_pag["columnas"]
        pag_rows = []
        for i, fila in enumerate(ct_pag["filas"]):
            row = [fila] + [
                str(ct_pag["frecuencias"][i][j]) for j in range(len(ct_pag["columnas"]))
            ]
            pag_rows.append(row)
        _add_table(doc, pag_headers, pag_rows)

    # Crosstab by date
    ct_fecha = data["crosstab_fecha"]
    if ct_fecha["columnas"] and ct_fecha["columnas"] != ["Sin fecha"]:
        _heading(doc, 2, "Tabla Cruzada: Categoría × Mes")

        fecha_headers = ["Categoría"] + ct_fecha["columnas"]
        fecha_rows = []
        for i, fila in enumerate(ct_fecha["filas"]):
            row = [fila] + [
                str(ct_fecha["frecuencias"][i][j]) for j in range(len(ct_fecha["columnas"]))
            ]
            fecha_rows.append(row)
        _add_table(doc, fecha_headers, fecha_rows)

    _heading(doc, 2, "Insights")

    if ct["alerta"]:
        _bullet(doc, ct["alerta"], "Patrón detectado: ")

    if ct_pag["alerta"]:
        _bullet(doc, ct_pag["alerta"], "Por página: ")

    if ct_fecha["alerta"]:
        _bullet(doc, ct_fecha["alerta"], "Temporal: ")


def _build_regla5(doc: Document, data: dict):
    kpis = data["kpis"]
    pie = data["pie_data"]
    bar = data["bar_data"]

    _heading(doc, 1, "5. Regla 5 — Dashboard (Visualización)")

    _para(
        doc,
        (
            "La quinta regla presenta los indicadores clave de forma visual. "
            "El dashboard incluye un gráfico de pastel (distribución violento/no violento), "
            "un gráfico de barras (distribución por categoría) y la tabla de frecuencias "
            "ubicada debajo de las visualizaciones."
        ),
    )

    _heading(doc, 2, "Distribución: Con/Sin Violencia")

    pie_rows = []
    total_pie = sum(p["Cantidad"] for p in pie)
    for p in pie:
        pct = round(p["Cantidad"] / total_pie * 100, 1) if total_pie else 0
        pie_rows.append([p["Estado"], str(p["Cantidad"]), f"{pct}%"])
    _add_table(doc, ["Estado", "Cantidad", "Porcentaje"], pie_rows, col_widths=[2, 1.5, 1.5])

    _heading(doc, 2, "Distribución por Categoría (Contenido Violento)")

    bar_rows = []
    for r in bar:
        bar_rows.append([r["Categoría"], str(r["Cantidad"]), f"{r['Porcentaje']}%"])
    _add_table(doc, ["Categoría", "Cantidad", "Porcentaje"], bar_rows, col_widths=[2.5, 1.2, 1.2])

    _heading(doc, 2, "Insights")

    violent_count = next((p["Cantidad"] for p in pie if p["Estado"] == "Con violencia"), 0)
    non_violent_count = next((p["Cantidad"] for p in pie if p["Estado"] == "Sin violencia"), 0)
    ratio = round(violent_count / non_violent_count, 2) if non_violent_count else 0

    _bullet(
        doc,
        (
            f"De {total_pie} registros válidos, {violent_count} ({kpis['violent_pct']}%) "
            f"contienen violencia de género y {non_violent_count} no. "
            f"La proporción violento/no violento es de 1:{ratio}."
        ),
        "Distribución general: ",
    )

    top_bar = bar[0] if bar else None
    if top_bar:
        _bullet(
            doc,
            (
                f"'{top_bar['Categoría']}' domina con {top_bar['Cantidad']} casos "
                f"({top_bar['Porcentaje']}% del contenido violento), seguida por "
                f"'{bar[1]['Categoría']}' ({bar[1]['Cantidad']} casos, {bar[1]['Porcentaje']}%)."
            ),
            "Categorías principales: ",
        )


def _build_regla6(doc: Document, data: dict):
    cm = data["confusion_matrix"]
    rm = data["reliability_metrics"]

    _heading(doc, 1, "6. Regla 6 — Métricas de Fiabilidad y Validez de la IA")

    _para(
        doc,
        (
            "La sexta regla evalúa la calidad del clasificador automático comparando "
            "sus predicciones contra el ground truth proporcionado por los revisores "
            "humanos. Se construye una matriz de confusión binaria (violencia vs. "
            "no violencia) y se calculan Precisión, Sensibilidad (Recall) y F1-Score."
        ),
    )

    _heading(doc, 2, "Matriz de Confusión")

    cm_headers = ["", "Real: Violencia", "Real: No Violencia"]
    cm_rows = [
        ["Predicho: Violencia", str(cm["VP"]), str(cm["FP"])],
        ["Predicho: No Violencia", str(cm["FN"]), str(cm["VN"])],
    ]
    _add_table(doc, cm_headers, cm_rows, col_widths=[2, 1.5, 1.5])

    _heading(doc, 2, "Métricas de Clasificación")

    met_rows = [
        ["Verdaderos Positivos (VP)", str(cm["VP"]), "IA dice violencia, humano confirma"],
        ["Verdaderos Negativos (VN)", str(cm["VN"]), "IA dice no violencia, humano confirma"],
        ["Falsos Positivos (FP)", str(cm["FP"]), "IA dice violencia, humano dice no"],
        ["Falsos Negativos (FN)", str(cm["FN"]), "IA dice no violencia, humano dice sí"],
        ["Total evaluado", str(cm["Total"]), "Soporte de la evaluación"],
    ]
    _add_table(doc, ["Métrica", "Valor", "Descripción"], met_rows, col_widths=[2.2, 1, 3])

    _heading(doc, 2, "Indicadores de Rendimiento")

    perf_rows = [
        [
            "Precisión",
            f"{rm['Precisión']:.4f}",
            f"{rm['Precisión'] * 100:.1f}%",
            "De los que la IA marca como violentos, cuántos realmente lo son",
        ],
        [
            "Sensibilidad (Recall)",
            f"{rm['Sensibilidad (Recall)']:.4f}",
            f"{rm['Sensibilidad (Recall)'] * 100:.1f}%",
            "De los violentos reales, cuántos detecta la IA",
        ],
        [
            "F1-Score",
            f"{rm['F1-Score']:.4f}",
            f"{rm['F1-Score'] * 100:.1f}%",
            "Media armónicaPrecisión y Sensibilidad",
        ],
        ["Soporte", str(rm["Soporte"]), "—", "Total de casos evaluados"],
    ]
    _add_table(
        doc,
        ["Indicador", "Valor", "Porcentaje", "Interpretación"],
        perf_rows,
        col_widths=[1.8, 1, 1, 3],
    )

    _heading(doc, 2, "Insights")

    precision_pct = rm["Precisión"] * 100
    recall_pct = rm["Sensibilidad (Recall)"] * 100
    f1_pct = rm["F1-Score"] * 100

    if precision_pct >= 80:
        _bullet(
            doc,
            (
                f"La precisión del {precision_pct:.1f}% indica que la IA es confiable "
                f"cuando clasifica contenido como violento — pocos falsos positivos."
            ),
            "Precisión: ",
        )
    else:
        _bullet(
            doc,
            (
                f"La precisión del {precision_pct:.1f}% indica que la IA genera "
                f"una cantidad significativa de falsos positivos ({cm['FP']} de "
                f"{cm['VP'] + cm['FP']}). Se recomienda calibrar el umbral de clasificación."
            ),
            "Precisión baja: ",
        )

    if recall_pct < 60:
        _bullet(
            doc,
            (
                f"La sensibilidad del {recall_pct:.1f}% es baja: la IA deja pasar "
                f"{cm['FN']} casos de violencia real ({cm['FN']} falsos negativos). "
                f"Esto es preocupante porque significa contenido violento no detectado."
            ),
            "Sensibilidad baja: ",
        )
    else:
        _bullet(
            doc,
            (
                f"La sensibilidad del {recall_pct:.1f}% es aceptable: la IA detecta "
                f"la mayoría de los casos violentos."
            ),
            "Sensibilidad aceptable: ",
        )

    _bullet(
        doc,
        (
            f"El F1-Score de {f1_pct:.1f}% representa el equilibrio entre precisión "
            f"y sensibilidad. "
            + (
                "Es un rendimiento razonable para un sistema de detección de contenido violento."
                if f1_pct >= 60
                else "Se recomienda mejorar el modelo o ajustar los prompts de clasificación."
            )
        ),
        "F1-Score: ",
    )


def _build_regla7(doc: Document, data: dict):
    val = data["validation_breakdown"]
    adj = data["adjustment_breakdown"]

    _heading(doc, 1, "7. Validación Humana (Human-in-the-Loop)")

    _para(
        doc,
        (
            "El sistema incorpora un ciclo de validación humana donde los revisores "
            "pueden estar o no de acuerdo con la clasificación automática. Las "
            "correcciones se almacenan en SQLite (fuente de verdad) y las "
            "correcciones con desacuerdo se indexan en ChromaDB como ejemplos "
            "few-shot para el clasificador RAG."
        ),
    )

    _heading(doc, 2, "Breakdown de Validación")

    val_rows = [
        ["Total analizable (neto)", str(val["net_total"])],
        ["Validados (acuerdo + desacuerdo)", f"{val['validated_count']} ({val['validated_pct']}%)"],
        ["  Acuerdo (AI = humano)", f"{val['agreed_count']}"],
        ["  Desacuerdo (humano corrige)", f"{val['disagreed_count']}"],
        ["Pendientes de revisión", f"{val['pending_count']} ({val['pending_pct']}%)"],
        ["Total original (con excluidos)", str(val["total"])],
    ]
    _add_table(doc, ["Métrica", "Valor"], val_rows, col_widths=[3, 2])

    _heading(doc, 2, "Breakdown de Ajustes")

    adj_rows = [
        [
            "Contenidos autónomos (AI correcta)",
            f"{adj['total'] - adj['adjusted_count']} ({adj['autonomous_pct']}%)",
        ],
        ["Contenidos ajustados por humano", f"{adj['adjusted_count']} ({adj['adjusted_pct']}%)"],
        ["Total", str(adj["total"])],
    ]
    _add_table(doc, ["Categoría", "Valor"], adj_rows, col_widths=[3, 2])

    _heading(doc, 2, "Insights")

    if val["validated_pct"] == 100:
        _bullet(
            doc,
            (
                "El 100% de los registros analizables han sido revisados por humanos. "
                "Esto proporciona un ground truth completo para evaluar al clasificador."
            ),
            "Cobertura completa: ",
        )
    else:
        _bullet(
            doc,
            (
                f"Solo el {val['validated_pct']}% de los registros han sido validados. "
                f"Quedan {val['pending_count']} registros pendientes de revisión."
            ),
            "Cobertura parcial: ",
        )

    agreement_rate = (
        round(val["agreed_count"] / val["validated_count"] * 100, 1)
        if val["validated_count"]
        else 0
    )
    _bullet(
        doc,
        (
            f"El {agreement_rate}% de las revisiones terminaron en acuerdo, lo que indica "
            f"{
                'una alta concordancia entre IA y humano'
                if agreement_rate >= 80
                else 'margen de mejora en la calibración del clasificador'
            }."
        ),
        "Tasa de acuerdo: ",
    )

    _bullet(
        doc,
        (
            f"El {adj['adjusted_pct']}% de los contenidos requirió corrección humana. "
            f"Estos {adj['adjusted_count']} casos son los que se indexan en ChromaDB "
            f"como ejemplos few-shot para mejorar futuras clasificaciones."
        ),
        "Intervención humana: ",
    )


def _build_conclusiones(doc: Document, data: dict):
    kpis = data["kpis"]
    rel = data["reliability"]
    rm = data["reliability_metrics"]
    val = data["validation_breakdown"]
    adj = data["adjustment_breakdown"]

    _heading(doc, 1, "8. Conclusiones y Recomendaciones")

    _para(doc, "Hallazgos clave del análisis estadístico:", bold=True)

    _bullet(
        doc,
        (
            f"El dataset contiene {data['meta']['total_analysis']} registros de los cuales "
            f"{kpis['violent_pct']}% son violentos. La categoría dominante es "
            f"'{kpis['top_category']}'."
        ),
    )

    _bullet(
        doc,
        (
            f"La fiabilidad de datos es CRÍTICA ({rel['pct_basura']}% de basura digital). "
            f"El scraper necesita revisión para reducir la captura de contenido vacío."
        ),
    )

    f1_msg = "rendimiento razonable" if rm["F1-Score"] >= 0.6 else "necesidad de calibración"
    _bullet(
        doc,
        (
            f"El clasificador RAG tiene una precisión del {rm['Precisión'] * 100:.1f}% y "
            f"una sensibilidad del {rm['Sensibilidad (Recall)'] * 100:.1f}%. "
            f"El F1-Score de {rm['F1-Score'] * 100:.1f}% indica {f1_msg}."
        ),
    )

    agr_count = val["agreed_count"]
    val_count = val["validated_count"]
    agr_pct = round(agr_count / val_count * 100, 1) if val_count else 0
    _bullet(
        doc,
        (
            f"El {val['validated_pct']}% del dataset ha sido validado por "
            f"humanos, con una tasa de acuerdo del {agr_pct}%. "
            f"El {adj['adjusted_pct']}% de contenidos requirió corrección."
        ),
    )

    _heading(doc, 2, "Recomendaciones")

    _bullet(
        doc,
        (
            "Revisar la configuración del scraper para filtrar contenido multimedia "
            "sin texto antes de persistirlo (COND_1_VACIO)."
        ),
        "Mejora del scraper: ",
    )

    _bullet(
        doc,
        (
            f"Calibrar el clasificador para reducir falsos negativos "
            f"({data['confusion_matrix']['FN']} casos) — priorizar recall sobre precisión "
            f"dado el dominio de aplicación (detección de violencia)."
        ),
        "Calibración del modelo: ",
    )

    _bullet(
        doc,
        (
            f"Indexar los {adj['adjusted_count']} desacuerdos en ChromaDB "
            "para enriquecer los ejemplos few-shot del RAGClassifier."
        ),
        "Retroalimentación RAG: ",
    )

    _bullet(
        doc,
        (
            "Continuar con el proceso de validación humana hasta alcanzar "
            "al menos el 80% de cobertura del dataset."
        ),
        "Cobertura de validación: ",
    )


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def generate_docx(stats_path: str, out_path: str):
    """Generate the full .docx report from a stats snapshot."""
    stats = json.loads(Path(stats_path).read_text())

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Build sections
    _build_portada(doc, stats["meta"])
    _build_resumen_ejecutivo(doc, stats)
    doc.add_page_break()

    _build_regla1(doc, stats)
    doc.add_page_break()

    _build_regla2(doc, stats)
    doc.add_page_break()

    _build_regla3(doc, stats)
    doc.add_page_break()

    _build_regla4(doc, stats)
    doc.add_page_break()

    _build_regla5(doc, stats)
    doc.add_page_break()

    _build_regla6(doc, stats)
    doc.add_page_break()

    _build_regla7(doc, stats)
    doc.add_page_break()

    _build_conclusiones(doc, stats)

    # Save
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Report saved to {out}")


def main():
    parser = argparse.ArgumentParser(description="Generate statistical report .docx")
    parser.add_argument("--stats", default=str(ROOT / "data" / "exports" / "stats_snapshot.json"))
    parser.add_argument("--out", default=str(ROOT / "docs" / "informe-estadistico-sistema.docx"))
    args = parser.parse_args()
    generate_docx(args.stats, args.out)


if __name__ == "__main__":
    main()
